"""
Resilient Document Parser - Safe Document Parsing with Fallbacks and Chunking.

Implements a robust document parser with:
- Multiple parser fallbacks for each file type
- Chunked processing for large files
- Memory-efficient streaming
- Comprehensive error handling
- Support for PDF, DOCX, TXT, MD, HTML, and more
"""

import logging
import asyncio
import io
import tempfile
import os
from typing import Any, Dict, List, Optional, Union, Tuple, BinaryIO, AsyncIterator
from pathlib import Path
from dataclasses import dataclass, field
from pydantic import BaseModel, Field
from enum import Enum
import json
import hashlib

logger = logging.getLogger(__name__)


class DocumentType(str, Enum):
    """Supported document types."""
    PDF = "pdf"
    DOCX = "docx"
    TXT = "txt"
    MD = "md"
    HTML = "html"
    RTF = "rtf"
    CSV = "csv"
    XLSX = "xlsx"
    PPTX = "pptx"
    EPUB = "epub"
    UNKNOWN = "unknown"


class ParserStatus(str, Enum):
    """Status of parsing operation."""
    SUCCESS = "success"
    FALLBACK_USED = "fallback_used"
    FAILED = "failed"
    PARTIAL = "partial"
    TIMEOUT = "timeout"


@dataclass
class ParseResult:
    """Result of document parsing."""
    content: str
    metadata: Dict[str, Any]
    status: ParserStatus
    parser_used: str
    chunks: List[str] = field(default_factory=list)
    error_message: Optional[str] = None
    warnings: List[str] = field(default_factory=list)
    processing_time_seconds: float = 0.0
    
    def add_warning(self, warning: str) -> None:
        """Add a warning message."""
        self.warnings.append(warning)


@dataclass
class DocumentChunk:
    """A chunk of document content."""
    content: str
    chunk_id: str
    start_page: Optional[int] = None
    end_page: Optional[int] = None
    metadata: Dict[str, Any] = field(default_factory=dict)


@dataclass
class ParserConfig:
    """Configuration for document parser."""
    # Chunking settings
    max_chunk_size_chars: int = 4000
    chunk_overlap_chars: int = 200
    respect_sentence_boundaries: bool = True
    
    # Memory limits
    max_file_size_mb: int = 100
    max_memory_usage_mb: int = 512
    
    # Timeout settings
    parse_timeout_seconds: float = 30.0
    chunk_timeout_seconds: float = 10.0
    
    # Parser preferences (order of fallback)
    pdf_parsers: List[str] = field(default_factory=lambda: ["pypdf2", "pdfplumber", "pymupdf"])
    docx_parsers: List[str] = field(default_factory=lambda: ["python-docx", "mammoth"])
    html_parsers: List[str] = field(default_factory=lambda: ["beautifulsoup", "lxml", "html5lib"])
    
    # Extraction options
    extract_images: bool = False
    extract_tables: bool = True
    extract_metadata: bool = True
    preserve_formatting: bool = False


class BaseDocumentParser:
    """Base class for document parsers."""
    
    def __init__(self, name: str):
        self.name = name
        self.logger = logging.getLogger(f"Parser.{name}")
    
    async def can_parse(self, file_path: Path, file_type: DocumentType) -> bool:
        """Check if this parser can handle the file."""
        raise NotImplementedError
    
    async def parse(self, file_path: Path, config: ParserConfig) -> ParseResult:
        """Parse the document."""
        raise NotImplementedError
    
    async def parse_stream(self, stream: BinaryIO, config: ParserConfig) -> ParseResult:
        """Parse from a stream."""
        # Default implementation writes to temp file
        with tempfile.NamedTemporaryFile(delete=False) as tmp:
            tmp.write(stream.read())
            tmp_path = Path(tmp.name)
        
        try:
            return await self.parse(tmp_path, config)
        finally:
            tmp_path.unlink(missing_ok=True)


class PDFParserPyPDF2(BaseDocumentParser):
    """PDF parser using PyPDF2."""
    
    def __init__(self):
        super().__init__("pypdf2")
    
    async def can_parse(self, file_path: Path, file_type: DocumentType) -> bool:
        return file_type == DocumentType.PDF
    
    async def parse(self, file_path: Path, config: ParserConfig) -> ParseResult:
        import PyPDF2
        
        result = ParseResult(
            content="",
            metadata={},
            status=ParserStatus.SUCCESS,
            parser_used=self.name
        )
        
        try:
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                
                # Extract metadata
                if config.extract_metadata:
                    result.metadata = {
                        "title": reader.metadata.get('/Title', ''),
                        "author": reader.metadata.get('/Author', ''),
                        "creator": reader.metadata.get('/Creator', ''),
                        "producer": reader.metadata.get('/Producer', ''),
                        "pages": len(reader.pages)
                    }
                
                # Extract text
                text_parts = []
                for page_num, page in enumerate(reader.pages):
                    try:
                        text = page.extract_text()
                        if text.strip():
                            text_parts.append(text)
                    except Exception as e:
                        result.add_warning(f"Failed to extract page {page_num}: {e}")
                
                result.content = "\n\n".join(text_parts)
                
        except Exception as e:
            result.status = ParserStatus.FAILED
            result.error_message = str(e)
        
        return result


class PDFParserPDFPlumber(BaseDocumentParser):
    """PDF parser using PDFPlumber."""
    
    def __init__(self):
        super().__init__("pdfplumber")
    
    async def can_parse(self, file_path: Path, file_type: DocumentType) -> bool:
        return file_type == DocumentType.PDF
    
    async def parse(self, file_path: Path, config: ParserConfig) -> ParseResult:
        import pdfplumber
        
        result = ParseResult(
            content="",
            metadata={},
            status=ParserStatus.SUCCESS,
            parser_used=self.name
        )
        
        try:
            with pdfplumber.open(file_path) as pdf:
                # Extract metadata
                if config.extract_metadata:
                    result.metadata = {
                        "title": pdf.metadata.get('Title', ''),
                        "author": pdf.metadata.get('Author', ''),
                        "creator": pdf.metadata.get('Creator', ''),
                        "producer": pdf.metadata.get('Producer', ''),
                        "pages": len(pdf.pages)
                    }
                
                # Extract text
                text_parts = []
                for page_num, page in enumerate(pdf.pages):
                    try:
                        text = page.extract_text()
                        if text.strip():
                            text_parts.append(text)
                            
                        # Extract tables if requested
                        if config.extract_tables:
                            tables = page.extract_tables()
                            if tables:
                                for table in tables:
                                    table_text = self._table_to_text(table)
                                    if table_text.strip():
                                        text_parts.append(table_text)
                        
                    except Exception as e:
                        result.add_warning(f"Failed to extract page {page_num}: {e}")
                
                result.content = "\n\n".join(text_parts)
                
        except Exception as e:
            result.status = ParserStatus.FAILED
            result.error_message = str(e)
        
        return result
    
    def _table_to_text(self, table: List[List[str]]) -> str:
        """Convert table to text format."""
        if not table:
            return ""
        
        # Simple table to text conversion
        rows = []
        for row in table:
            if row:
                rows.append(" | ".join(str(cell) if cell else "" for cell in row))
        
        return "\n".join(rows)


class PDFParserPyMuPDF(BaseDocumentParser):
    """PDF parser using PyMuPDF (fitz)."""
    
    def __init__(self):
        super().__init__("pymupdf")
    
    async def can_parse(self, file_path: Path, file_type: DocumentType) -> bool:
        return file_type == DocumentType.PDF
    
    async def parse(self, file_path: Path, config: ParserConfig) -> ParseResult:
        import fitz
        
        result = ParseResult(
            content="",
            metadata={},
            status=ParserStatus.SUCCESS,
            parser_used=self.name
        )
        
        try:
            doc = fitz.open(file_path)
            
            # Extract metadata
            if config.extract_metadata:
                result.metadata = {
                    "title": doc.metadata.get('title', ''),
                    "author": doc.metadata.get('author', ''),
                    "creator": doc.metadata.get('creator', ''),
                    "producer": doc.metadata.get('producer', ''),
                    "pages": doc.page_count
                }
            
            # Extract text
            text_parts = []
            for page_num in range(doc.page_count):
                try:
                    page = doc[page_num]
                    text = page.get_text()
                    if text.strip():
                        text_parts.append(text)
                except Exception as e:
                    result.add_warning(f"Failed to extract page {page_num}: {e}")
            
            result.content = "\n\n".join(text_parts)
            doc.close()
            
        except Exception as e:
            result.status = ParserStatus.FAILED
            result.error_message = str(e)
        
        return result


class DOCXParserPythonDocx(BaseDocumentParser):
    """DOCX parser using python-docx."""
    
    def __init__(self):
        super().__init__("python-docx")
    
    async def can_parse(self, file_path: Path, file_type: DocumentType) -> bool:
        return file_type == DocumentType.DOCX
    
    async def parse(self, file_path: Path, config: ParserConfig) -> ParseResult:
        from docx import Document
        
        result = ParseResult(
            content="",
            metadata={},
            status=ParserStatus.SUCCESS,
            parser_used=self.name
        )
        
        try:
            doc = Document(file_path)
            
            # Extract metadata
            if config.extract_metadata:
                core_props = doc.core_properties
                result.metadata = {
                    "title": core_props.title or '',
                    "author": core_props.author or '',
                    "created": str(core_props.created) if core_props.created else '',
                    "modified": str(core_props.modified) if core_props.modified else '',
                    "paragraphs": len(doc.paragraphs)
                }
            
            # Extract text
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract tables if requested
            if config.extract_tables:
                for table in doc.tables:
                    table_text = self._table_to_text(table)
                    if table_text.strip():
                        text_parts.append(table_text)
            
            result.content = "\n\n".join(text_parts)
            
        except Exception as e:
            result.status = ParserStatus.FAILED
            result.error_message = str(e)
        
        return result
    
    def _table_to_text(self, table) -> str:
        """Convert table to text format."""
        rows = []
        for row in table.rows:
            cells = [cell.text for cell in row.cells]
            rows.append(" | ".join(cells))
        return "\n".join(rows)


class TextParser(BaseDocumentParser):
    """Parser for plain text files."""
    
    def __init__(self):
        super().__init__("text")
    
    async def can_parse(self, file_path: Path, file_type: DocumentType) -> bool:
        return file_type in [DocumentType.TXT, DocumentType.MD, DocumentType.RTF]
    
    async def parse(self, file_path: Path, config: ParserConfig) -> ParseResult:
        result = ParseResult(
            content="",
            metadata={},
            status=ParserStatus.SUCCESS,
            parser_used=self.name
        )
        
        try:
            # Determine encoding
            encoding = 'utf-8'
            try:
                with open(file_path, 'rb') as f:
                    raw = f.read(10000)
                    if b'\0' in raw:
                        encoding = 'utf-16'
            except:
                pass
            
            # Read file
            with open(file_path, 'r', encoding=encoding) as f:
                content = f.read()
            
            result.content = content
            
            # Extract metadata
            if config.extract_metadata:
                stat = file_path.stat()
                result.metadata = {
                    "size_bytes": stat.st_size,
                    "encoding": encoding,
                    "lines": len(content.splitlines()),
                    "characters": len(content)
                }
            
        except Exception as e:
            result.status = ParserStatus.FAILED
            result.error_message = str(e)
        
        return result


class HTMLParser(BaseDocumentParser):
    """HTML parser using BeautifulSoup."""
    
    def __init__(self):
        super().__init__("beautifulsoup")
    
    async def can_parse(self, file_path: Path, file_type: DocumentType) -> bool:
        return file_type == DocumentType.HTML
    
    async def parse(self, file_path: Path, config: ParserConfig) -> ParseResult:
        from bs4 import BeautifulSoup
        
        result = ParseResult(
            content="",
            metadata={},
            status=ParserStatus.SUCCESS,
            parser_used=self.name
        )
        
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                html = f.read()
            
            soup = BeautifulSoup(html, 'html.parser')
            
            # Extract text
            if config.preserve_formatting:
                # Preserve some formatting
                text = soup.get_text(separator='\n', strip=True)
            else:
                # Clean text
                for script in soup(["script", "style"]):
                    script.decompose()
                text = soup.get_text(separator=' ', strip=True)
            
            result.content = text
            
            # Extract metadata
            if config.extract_metadata:
                title = soup.find('title')
                result.metadata = {
                    "title": title.get_text() if title else '',
                    "links": len(soup.find_all('a')),
                    "images": len(soup.find_all('img')),
                    "forms": len(soup.find_all('form'))
                }
            
        except Exception as e:
            result.status = ParserStatus.FAILED
            result.error_message = str(e)
        
        return result


class ResilientDocumentParser:
    """
    Resilient document parser with fallbacks and chunking.
    
    Features:
    - Multiple parser fallbacks for each file type
    - Chunked processing for large files
    - Memory-efficient streaming
    - Comprehensive error handling
    """
    
    def __init__(self, config: Optional[ParserConfig] = None):
        """Initialize resilient document parser.
        
        Args:
            config: Parser configuration
        """
        self.config = config or ParserConfig()
        
        # Initialize parsers
        self.parsers = {
            DocumentType.PDF: [
                PDFParserPyPDF2(),
                PDFParserPDFPlumber(),
                PDFParserPyMuPDF()
            ],
            DocumentType.DOCX: [
                DOCXParserPythonDocx()
            ],
            DocumentType.TXT: [
                TextParser()
            ],
            DocumentType.MD: [
                TextParser()
            ],
            DocumentType.HTML: [
                HTMLParser()
            ]
        }
        
        # Statistics
        self.stats = {
            "total_parsers": 0,
            "successful_parses": 0,
            "fallback_used": 0,
            "failed_parses": 0
        }
        
        # Count total parsers
        for parsers in self.parsers.values():
            self.stats["total_parsers"] += len(parsers)
    
    async def parse_document(
        self,
        file_path: Union[str, Path],
        chunk: bool = False
    ) -> ParseResult:
        """Parse a document with fallbacks.
        
        Args:
            file_path: Path to the document
            chunk: Whether to chunk the content
            
        Returns:
            ParseResult with parsed content
        """
        file_path = Path(file_path)
        
        # Validate file
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        # Check file size
        file_size_mb = file_path.stat().st_size / (1024 * 1024)
        if file_size_mb > self.config.max_file_size_mb:
            raise ValueError(f"File too large: {file_size_mb:.1f}MB > {self.config.max_file_size_mb}MB")
        
        # Determine file type
        file_type = self._detect_file_type(file_path)
        
        # Get parsers for this type
        parsers = self.parsers.get(file_type, [])
        if not parsers:
            raise ValueError(f"No parsers available for file type: {file_type}")
        
        # Try parsers in order
        last_error = None
        for parser in parsers:
            try:
                self.logger.info(f"Attempting to parse {file_path.name} with {parser.name}")
                
                # Parse with timeout
                result = await asyncio.wait_for(
                    parser.parse(file_path, self.config),
                    timeout=self.config.parse_timeout_seconds
                )
                
                if result.status == ParserStatus.SUCCESS:
                    self.stats["successful_parses"] += 1
                    
                    # Chunk content if requested
                    if chunk and result.content:
                        result.chunks = self._chunk_content(result.content)
                    
                    return result
                else:
                    last_error = result.error_message
                    
            except asyncio.TimeoutError:
                last_error = f"Parser {parser.name} timed out"
                self.logger.warning(last_error)
                continue
            except Exception as e:
                last_error = str(e)
                self.logger.warning(f"Parser {parser.name} failed: {e}")
                continue
        
        # All parsers failed
        self.stats["failed_parses"] += 1
        result = ParseResult(
            content="",
            metadata={},
            status=ParserStatus.FAILED,
            parser_used="none",
            error_message=last_error or "All parsers failed"
        )
        
        return result
    
    async def parse_stream(
        self,
        stream: BinaryIO,
        file_type: DocumentType,
        filename: str = "stream",
        chunk: bool = False
    ) -> ParseResult:
        """Parse from a stream with fallbacks.
        
        Args:
            stream: Binary stream
            file_type: Type of document
            filename: Original filename (for metadata)
            chunk: Whether to chunk the content
            
        Returns:
            ParseResult with parsed content
        """
        # Get parsers for this type
        parsers = self.parsers.get(file_type, [])
        if not parsers:
            raise ValueError(f"No parsers available for file type: {file_type}")
        
        # Try parsers in order
        last_error = None
        for parser in parsers:
            try:
                self.logger.info(f"Attempting to parse stream with {parser.name}")
                
                # Parse with timeout
                result = await asyncio.wait_for(
                    parser.parse_stream(stream, self.config),
                    timeout=self.config.parse_timeout_seconds
                )
                
                if result.status == ParserStatus.SUCCESS:
                    self.stats["successful_parses"] += 1
                    
                    # Update metadata with filename
                    result.metadata["filename"] = filename
                    
                    # Chunk content if requested
                    if chunk and result.content:
                        result.chunks = self._chunk_content(result.content)
                    
                    return result
                else:
                    last_error = result.error_message
                    
            except asyncio.TimeoutError:
                last_error = f"Parser {parser.name} timed out"
                self.logger.warning(last_error)
                continue
            except Exception as e:
                last_error = str(e)
                self.logger.warning(f"Parser {parser.name} failed: {e}")
                continue
        
        # All parsers failed
        self.stats["failed_parses"] += 1
        result = ParseResult(
            content="",
            metadata={"filename": filename},
            status=ParserStatus.FAILED,
            parser_used="none",
            error_message=last_error or "All parsers failed"
        )
        
        return result
    
    async def parse_large_document(
        self,
        file_path: Union[str, Path],
        chunk_size: Optional[int] = None
    ) -> AsyncIterator[DocumentChunk]:
        """Parse a large document in chunks.
        
        Args:
            file_path: Path to the document
            chunk_size: Override default chunk size
            
        Yields:
            DocumentChunk instances
        """
        file_path = Path(file_path)
        file_type = self._detect_file_type(file_path)
        
        # For now, we'll parse the whole document and chunk it
        # In a real implementation, you might want streaming parsers
        result = await self.parse_document(file_path, chunk=False)
        
        if result.status != ParserStatus.SUCCESS:
            yield DocumentChunk(
                content="",
                chunk_id="error",
                metadata={"error": result.error_message}
            )
            return
        
        # Create chunks
        chunk_size = chunk_size or self.config.max_chunk_size_chars
        chunks = self._chunk_content(result.content, chunk_size)
        
        for i, chunk_content in enumerate(chunks):
            chunk = DocumentChunk(
                content=chunk_content,
                chunk_id=f"{file_path.stem}_{i}",
                metadata={
                    "source_file": str(file_path),
                    "chunk_index": i,
                    "total_chunks": len(chunks)
                }
            )
            yield chunk
    
    def _detect_file_type(self, file_path: Path) -> DocumentType:
        """Detect document type from file extension."""
        ext = file_path.suffix.lower()
        
        type_map = {
            '.pdf': DocumentType.PDF,
            '.docx': DocumentType.DOCX,
            '.txt': DocumentType.TXT,
            '.md': DocumentType.MD,
            '.html': DocumentType.HTML,
            '.htm': DocumentType.HTML,
            '.rtf': DocumentType.RTF,
            '.csv': DocumentType.CSV,
            '.xlsx': DocumentType.XLSX,
            '.pptx': DocumentType.PPTX,
            '.epub': DocumentType.EPUB
        }
        
        return type_map.get(ext, DocumentType.UNKNOWN)
    
    def _chunk_content(
        self,
        content: str,
        chunk_size: Optional[int] = None
    ) -> List[str]:
        """Chunk content into smaller pieces.
        
        Args:
            content: Text to chunk
            chunk_size: Size of each chunk
            
        Returns:
            List of text chunks
        """
        if not content:
            return []
        
        chunk_size = chunk_size or self.config.max_chunk_size_chars
        overlap = self.config.chunk_overlap_chars
        
        if self.config.respect_sentence_boundaries:
            return self._chunk_by_sentences(content, chunk_size, overlap)
        else:
            return self._chunk_by_chars(content, chunk_size, overlap)
    
    def _chunk_by_chars(self, content: str, chunk_size: int, overlap: int) -> List[str]:
        """Chunk content by characters with overlap."""
        chunks = []
        start = 0
        
        while start < len(content):
            end = start + chunk_size
            chunk = content[start:end]
            
            if chunk.strip():
                chunks.append(chunk)
            
            start = end - overlap if end < len(content) else len(content)
        
        return chunks
    
    def _chunk_by_sentences(self, content: str, chunk_size: int, overlap: int) -> List[str]:
        """Chunk content respecting sentence boundaries."""
        import re
        
        # Split into sentences
        sentences = re.split(r'(?<=[.!?])\s+', content)
        
        chunks = []
        current_chunk = ""
        current_size = 0
        
        for sentence in sentences:
            sentence_size = len(sentence)
            
            # If adding this sentence would exceed chunk size
            if current_size + sentence_size > chunk_size and current_chunk:
                chunks.append(current_chunk.strip())
                
                # Start new chunk with overlap
                if overlap > 0 and chunks:
                    # Find sentences to overlap
                    overlap_text = chunks[-1][-overlap:] if len(chunks[-1]) > overlap else chunks[-1]
                    current_chunk = overlap_text + " " + sentence
                    current_size = len(current_chunk)
                else:
                    current_chunk = sentence
                    current_size = sentence_size
            else:
                current_chunk += " " + sentence if current_chunk else sentence
                current_size += sentence_size
        
        # Add final chunk
        if current_chunk.strip():
            chunks.append(current_chunk.strip())
        
        return chunks
    
    def get_stats(self) -> Dict[str, Any]:
        """Get parsing statistics."""
        total = self.stats["successful_parses"] + self.stats["failed_parses"]
        success_rate = self.stats["successful_parses"] / total if total > 0 else 0
        
        return {
            **self.stats,
            "success_rate": success_rate,
            "fallback_rate": self.stats["fallback_used"] / total if total > 0 else 0
        }
    
    def add_parser(self, file_type: DocumentType, parser: BaseDocumentParser) -> None:
        """Add a new parser for a file type.
        
        Args:
            file_type: Document type
            parser: Parser instance
        """
        if file_type not in self.parsers:
            self.parsers[file_type] = []
        
        self.parsers[file_type].append(parser)
        self.stats["total_parsers"] += 1


# Factory function for creating resilient document parser
def create_resilient_document_parser(config: Optional[ParserConfig] = None) -> ResilientDocumentParser:
    """Create a resilient document parser.
    
    Args:
        config: Parser configuration
        
    Returns:
        ResilientDocumentParser instance
    """
    return ResilientDocumentParser(config)

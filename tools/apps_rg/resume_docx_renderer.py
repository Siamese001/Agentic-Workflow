"""Render a generated_resume.json into a .docx matching SVP Engineering Resume_Ayer.docx style.

Usage:
    python tools/apps_rg/resume_docx_renderer.py <generated_resume.json> <output.docx> [--template <template.docx>]

Copies styles from the template docx, then writes:
  Heading 1   — candidate name
  Normal       — title tagline + contact line
  Heading 2   — section headers (EXECUTIVE SUMMARY, PROFESSIONAL EXPERIENCE, etc.)
  Normal       — company/role/dates lines
  Normal/Body  — achievement bullets (prefixed with bold keyword label)
  Heading 3   — education + certification lines
  Normal       — skills competency lines
"""

from __future__ import annotations

import argparse
import json
import sys
from pathlib import Path

import docx
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
import copy


_DEFAULT_TEMPLATE = Path(r"C:\Users\amita\Documents\Resumes\SVP Engineering Resume_Ayer.docx")
_DEFAULT_NAME = "Amit Ayer"
_DEFAULT_TAGLINE = "SVP IT Strategy & Innovation | Agentic AI Platforms | Enterprise Architecture | AI/ML"
_DEFAULT_CONTACT = "+1-917-239-3830 | amitayer1@gmail.com | linkedin.com/in/amitayer1 | github.com/Siamese001"


def _add(doc: Document, style: str, text: str, bold_prefix: str | None = None) -> None:
    """Add a paragraph with given style. Optionally bold the first segment up to ': '."""
    p = doc.add_paragraph(style=style)
    if bold_prefix and ": " in text:
        label, rest = text.split(": ", 1)
        run = p.add_run(label + ": ")
        run.bold = True
        p.add_run(rest)
    else:
        p.add_run(text)


def render(resume_json: dict, output_path: Path, template_path: Path) -> None:
    doc = Document(template_path)

    # Clear all existing content
    for el in list(doc.element.body):
        tag = el.tag.split("}")[-1] if "}" in el.tag else el.tag
        if tag in ("p", "tbl", "sdt"):
            doc.element.body.remove(el)

    name = resume_json.get("name", _DEFAULT_NAME)
    tagline = resume_json.get("tagline", _DEFAULT_TAGLINE)
    contact = resume_json.get("contact", _DEFAULT_CONTACT)

    # Header block
    _add(doc, "Heading 1", name)
    _add(doc, "Normal", tagline)
    _add(doc, "Normal", contact)
    doc.add_paragraph(style="Body Text")  # spacer

    # Executive Summary
    exec_summary = resume_json.get("executive_summary", {})
    if isinstance(exec_summary, dict):
        raw = exec_summary.get("content", [])
        # Qwen may split into array-of-sentences; join into one paragraph
        if isinstance(raw, list):
            content_lines = [" ".join(str(s).strip() for s in raw if str(s).strip())]
        else:
            content_lines = [str(raw)] if raw else []
    elif isinstance(exec_summary, str):
        content_lines = [exec_summary]
    else:
        content_lines = []

    _add(doc, "Heading 2", "EXECUTIVE SUMMARY")
    for line in content_lines:
        _add(doc, "Body Text", line)
    doc.add_paragraph(style="Body Text")

    # Professional Experience
    experience = resume_json.get("experience", [])
    if experience:
        _add(doc, "Heading 2", "PROFESSIONAL EXPERIENCE")
        for role in experience:
            company = role.get("company", "")
            location = role.get("location", "")
            title = role.get("title", role.get("role", ""))
            dates = role.get("dates", "")

            company_line = f"{company} | {location}" if location else company
            _add(doc, "Normal", company_line)
            _add(doc, "Normal", f"{title} | {dates}" if dates else title)

            # Intro sentence — try intro, summary, or first item of description list
            intro = role.get("intro") or role.get("summary", "")
            desc_list = role.get("description", [])
            if not intro and isinstance(desc_list, list) and desc_list:
                intro = desc_list[0]
                desc_list = desc_list[1:]  # remaining become bullets
            elif not intro and isinstance(desc_list, str):
                intro = desc_list
                desc_list = []
            if intro:
                _add(doc, "Normal", intro)

            # Bullets — try bullets key first, then fall back to remaining desc_list
            bullets = role.get("bullets", [])
            if not bullets and desc_list:
                bullets = [{"description": b} for b in desc_list]
            for b in bullets:
                if isinstance(b, dict):
                    desc = b.get("description", "")
                    anchor = b.get("evidence_anchor", "")
                    # If desc already starts with anchor label, don't double it
                    text = desc if (anchor and desc.startswith(anchor)) else (f"{anchor}: {desc}" if anchor else desc)
                else:
                    text = str(b)
                _add(doc, "Normal", text, bold_prefix=True)

            doc.add_paragraph(style="Normal")  # spacer between roles

    # Education
    education = resume_json.get("education", [])
    if education:
        _add(doc, "Heading 2", "EDUCATION")
        for edu in education:
            if isinstance(edu, dict):
                degree = edu.get("degree", "")
                institution = edu.get("institution", "")
                year = edu.get("graduation_year", edu.get("year", ""))
                parts = [p for p in [degree, institution, year] if p]
                _add(doc, "Heading 3", ", ".join(parts))
            else:
                _add(doc, "Heading 3", str(edu))
        doc.add_paragraph(style="Body Text")

    # Certifications
    certs = resume_json.get("certifications", [])
    if certs:
        _add(doc, "Heading 2", "CERTIFICATIONS & CREDENTIALS")
        for cert in certs:
            if isinstance(cert, dict):
                parts = [p for p in [cert.get("name", ""), cert.get("issuer", ""), str(cert.get("year", ""))] if p]
                _add(doc, "Heading 3", ", ".join(parts))
            else:
                _add(doc, "Heading 3", str(cert))
        doc.add_paragraph(style="Body Text")

    # Skills
    skills = resume_json.get("skills", [])
    competencies = resume_json.get("competencies", resume_json.get("engineering_competencies", []))
    if skills or competencies:
        _add(doc, "Heading 2", "ENGINEERING & PLATFORM COMPETENCIES")
        for skill in skills:
            if isinstance(skill, dict):
                label = skill.get("label", skill.get("name", ""))
                detail = skill.get("detail", skill.get("description", ""))
                text = f"{label}: {detail}" if detail else label
                _add(doc, "Normal", text, bold_prefix=True)
            else:
                _add(doc, "Normal", str(skill))
        for comp in competencies:
            if isinstance(comp, dict):
                label = comp.get("label", comp.get("name", ""))
                detail = comp.get("detail", comp.get("description", ""))
                text = f"{label}: {detail}" if detail else label
                _add(doc, "Normal", text, bold_prefix=True)
            else:
                _add(doc, "Normal", str(comp))

    doc.save(output_path)
    print(f"Saved: {output_path}")


def main() -> None:
    parser = argparse.ArgumentParser(description="Render generated_resume.json to docx")
    parser.add_argument("resume_json", help="Path to generated_resume.json")
    parser.add_argument("output_docx", help="Output .docx path")
    parser.add_argument("--template", default=str(_DEFAULT_TEMPLATE), help="Template .docx")
    args = parser.parse_args()

    with open(args.resume_json, encoding="utf-8") as f:
        data = json.load(f)

    render(data, Path(args.output_docx), Path(args.template))


if __name__ == "__main__":
    main()

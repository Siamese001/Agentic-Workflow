"""Render apps_rg generated_resume.json into a .docx using the source resume as style template.

Usage:
    python tools/apps_rg/render_resume_docx.py \
        --template path/to/source_resume.docx \
        --json artifacts/apps_rg/runs/<run>/generated_resume.json \
        --output artifacts/apps_rg/runs/<run>/generated_resume.docx

The script clones the template's styles, margins, and section properties, then
writes the generated content using the same style names found in the template:
    - Heading 1: candidate name
    - Normal (12pt bold): subtitle / company headers
    - Normal (10pt): contact line
    - Heading 2: section headers (EXECUTIVE SUMMARY, PROFESSIONAL EXPERIENCE, etc.)
    - Heading 3: certifications, education entries
    - Normal (bold run prefix): bullet items with bold label prefix
    - Body Text: body paragraphs
"""
from __future__ import annotations

import argparse
import json
import sys
from pathlib import Path

import docx
from docx.shared import Pt, Emu


def _clear_document(doc: docx.Document) -> None:
    """Remove all paragraphs and tables from the document body."""
    body = doc.element.body
    for child in list(body):
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag in ("p", "tbl"):
            body.remove(child)


def _add_heading(doc: docx.Document, text: str, level: int) -> None:
    """Add a heading paragraph using the template's heading style."""
    style_name = f"Heading {level}"
    p = doc.add_paragraph(text, style=style_name)
    return p


def _add_normal_bold(doc: docx.Document, text: str, size: int | None = None) -> None:
    """Add a Normal-style paragraph with bold text and optional explicit size."""
    p = doc.add_paragraph(style="Normal")
    run = p.add_run(text)
    run.bold = True
    if size:
        run.font.size = Pt(size)


def _add_normal(doc: docx.Document, text: str, size: int | None = None) -> None:
    """Add a Normal-style paragraph."""
    p = doc.add_paragraph(style="Normal")
    run = p.add_run(text)
    if size:
        run.font.size = Pt(size)


def _add_bullet_with_label(doc: docx.Document, label: str, description: str) -> None:
    """Add a bullet item: bold label prefix followed by normal description text."""
    p = doc.add_paragraph(style="Normal")
    if label:
        run_label = p.add_run(f"{label}: ")
        run_label.bold = True
    run_desc = p.add_run(description)


def _add_body(doc: docx.Document, text: str) -> None:
    """Add a Body Text paragraph."""
    p = doc.add_paragraph(text, style="Body Text")


def render(template_path: str, json_path: str, output_path: str) -> Path:
    """Render the generated resume JSON into a formatted .docx."""
    with open(json_path, "r", encoding="utf-8") as f:
        data = json.load(f)

    # Open template to inherit styles and page setup
    doc = docx.Document(template_path)
    _clear_document(doc)

    # --- NAME ---
    name = "Amit Ayer"  # Preserved from source resume
    _add_heading(doc, name, level=1)

    # --- SUBTITLE (role line from target) ---
    target_role = data.get("target_role", "")
    skills_short = data.get("skills", [])[:4]
    subtitle_parts = [target_role] + skills_short if target_role else skills_short
    subtitle = " | ".join(subtitle_parts)
    _add_normal_bold(doc, subtitle, size=12)

    # --- CONTACT ---
    contact = "+1-917-239-3830 | amitayer1@gmail.com | linkedin.com/in/amitayer1 | github.com/Siamese001"
    _add_normal(doc, contact, size=10)

    # --- EXECUTIVE SUMMARY ---
    _add_heading(doc, "EXECUTIVE SUMMARY", level=2)
    summary = data.get("executive_summary", "")
    if isinstance(summary, dict):
        summary = summary.get("summary", "")
    _add_body(doc, summary)

    # --- PROFESSIONAL EXPERIENCE ---
    _add_heading(doc, "PROFESSIONAL EXPERIENCE", level=2)
    for exp in data.get("experience", []):
        company = exp.get("company", "")
        location = exp.get("location", "")
        title = exp.get("title", "")
        dates = exp.get("dates", "")

        _add_normal_bold(doc, f"{company} | {location}", size=12)
        _add_normal_bold(doc, f"{title} | {dates}")

        # Lead-in paragraph (if present)
        lead_in = exp.get("lead_in", "")
        if lead_in:
            _add_normal(doc, lead_in)

        # Bullets
        for bullet in exp.get("bullets", []):
            if isinstance(bullet, dict):
                desc = bullet.get("description", "")
                # Some bullets have bold prefix labels
                _add_bullet_with_label(doc, "", desc)
            else:
                _add_bullet_with_label(doc, "", str(bullet))

    # --- EDUCATION ---
    education = data.get("education", [])
    if education:
        _add_heading(doc, "EDUCATION", level=2)
        for edu in education:
            if isinstance(edu, dict):
                degree = edu.get("degree", "")
                institution = edu.get("institution", "")
                year = edu.get("graduation_year", "")
                line = f"{degree}, {institution}"
                if year:
                    line += f" ({year})"
                _add_heading(doc, line, level=3)
            else:
                _add_heading(doc, str(edu), level=3)

    # --- CERTIFICATIONS ---
    certs = data.get("certifications", [])
    if certs:
        _add_heading(doc, "CERTIFICATIONS & CREDENTIALS", level=2)
        for cert in certs:
            if isinstance(cert, dict):
                name_c = cert.get("name", "")
                issuer = cert.get("issuer", "")
                year = cert.get("year", "")
                parts = [name_c]
                if issuer:
                    parts.append(issuer)
                if year:
                    parts.append(year)
                _add_heading(doc, ", ".join(parts), level=3)
            else:
                _add_heading(doc, str(cert), level=3)

    # --- SKILLS / COMPETENCIES ---
    skills = data.get("skills", [])
    if skills:
        _add_heading(doc, "ENGINEERING & PLATFORM COMPETENCIES", level=2)
        for skill in skills:
            _add_normal_bold(doc, skill)

    # Save
    out = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    return out


def main() -> None:
    parser = argparse.ArgumentParser(description="Render apps_rg JSON to .docx")
    parser.add_argument("--template", required=True, help="Path to source resume .docx (style template)")
    parser.add_argument("--json", required=True, help="Path to generated_resume.json")
    parser.add_argument("--output", required=True, help="Output .docx path")
    args = parser.parse_args()

    out = render(args.template, args.json, args.output)
    print(f"Resume rendered: {out}")


if __name__ == "__main__":
    main()

"""Deterministic DOCX emission from assembly final_resume.json + DOCX manifest (no providers).

Rules:
- Canonical text payloads are sliced **only** from final_resume structured fields.
- Locked sections emit **verbatim copied_text_exact** strings (typically JSON payloads in one paragraph).
"""

from __future__ import annotations

if __name__ == "__main__":
    raise ImportError(
        "This module is not an operator CLI entrypoint. "
        "Use: python -m apps_rg or python -m apps_rg --section <lane>"
    )

from dataclasses import dataclass
from datetime import datetime, timezone
import hashlib
import json
from pathlib import Path
from typing import Any

from apps_rg.runtime.locked_copy.locked_copy_manifest import find_repo_root
from apps_rg.runtime.render.docx_render_x2 import (
    failures,
    gate_records_to_blob,
    gates_all_pass,
    run_docx_render_x2_gates,
)

_EXPECTED_DOCX_BASENAME = "amit_ayer_resume_v1.docx"


try:
    from docx import Document  # type: ignore[import-untyped]
except ImportError:
    Document = None  # type: ignore[misc, assignment]


def _sha256_file(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def _heading_style_name(level: int) -> str:
    clipped = max(1, min(int(level), 9))
    return f"Heading {clipped}"


def _paragraph_safe(doc: Any, text: str, *, style_name: str | None) -> None:
    if style_name:
        try:
            doc.add_paragraph(text, style=style_name)
            return
        except (KeyError, ValueError):
            pass
    doc.add_paragraph(text)


def extract_plaintext_blocks(*, sid: str, section_blob: dict[str, Any]) -> list[str]:
    kind = section_blob.get("section_kind")
    if kind == "locked_copy_inline":
        ct = section_blob.get("copied_text_exact")
        return [] if ct is None else [str(ct)]
    if kind != "generated_lane":
        return []
    l2 = section_blob.get("l2_output_snapshot")
    if not isinstance(l2, dict):
        return []

    out: list[str] = []

    match sid:
        case "headline":
            line = l2.get("headline_line")
            if isinstance(line, str) and line:
                out.append(line)
        case "executive_summary":
            rdt = l2.get("resume_display_text")
            if isinstance(rdt, str) and rdt:
                out.append(rdt)
            else:
                raise ValueError(f"{sid}: missing resume_display_text for render")
        case "unify_narrative":
            ns = l2.get("narrative_sentence")
            if isinstance(ns, str) and ns:
                out.append(ns)
            else:
                raise ValueError(f"{sid}: missing narrative_sentence")
        case "ibm_narrative":
            ns = l2.get("narrative_sentence")
            if isinstance(ns, str) and ns:
                out.append(ns)
            else:
                raise ValueError(f"{sid}: missing narrative_sentence")
        case "unify_bullets" | "ibm_bullets":
            for b in l2.get("bullets") or []:
                if isinstance(b, dict):
                    bt = b.get("bullet_text")
                    if bt is None:
                        continue
                    bt_s = str(bt).strip()
                    if bt_s:
                        out.append(bt_s)
                else:
                    continue
        case "competencies":
            for cat in l2.get("competencies") or []:
                if not isinstance(cat, dict):
                    continue
                lab = str(cat.get("category_label") or "").strip()
                if lab:
                    out.append(lab)
                for term in cat.get("terms") or []:
                    ts = str(term).strip()
                    if ts:
                        out.append(ts)
        case _:
            raise ValueError(f"Unhandled generated lane for extraction: {sid}")

    cleaned = [x for x in out if isinstance(x, str) and x.strip()]
    if sid == "headline" and not cleaned:
        raise ValueError(f"{sid}: empty headline_line after trim")
    if sid == "executive_summary" and not cleaned:
        raise ValueError(f"{sid}: missing resume_display_text content")
    if sid not in ("headline", "executive_summary") and not cleaned:
        raise ValueError(f"{sid}: empty extracted plaintext payload")
    return cleaned


def _render_competency_section(doc: Any, section_blob: dict[str, Any]) -> None:
    l2 = section_blob.get("l2_output_snapshot")
    if not isinstance(l2, dict):
        return
    cats = [c for c in (l2.get("competencies") or []) if isinstance(c, dict)]
    for cat in cats:
        lab = str(cat.get("category_label") or "").strip()
        if not lab:
            continue
        _paragraph_safe(doc, lab, style_name="Heading 4")
        for term in cat.get("terms") or []:
            ts = str(term).strip()
            if not ts:
                continue
            _paragraph_safe(doc, ts, style_name="List Bullet")


@dataclass(frozen=True)
class DocxRendererPaths:
    repo_root: Path
    final_resume_json: Path
    docx_manifest_json: Path
    output_dir: Path
    output_docx: Path

    def rel(self, p: Path) -> str:
        try:
            return p.relative_to(self.repo_root).as_posix()
        except ValueError:
            return p.resolve().as_posix()


def resolve_docx_renderer_paths(repo: Path | None = None) -> DocxRendererPaths:
    root = repo or find_repo_root()
    out = root / "artifacts/apps_rg/runtime_proofs/docx"
    return DocxRendererPaths(
        repo_root=root,
        final_resume_json=root / "artifacts/apps_rg/runtime_proofs/final_resume_assembly/final_resume.json",
        docx_manifest_json=root / "artifacts/apps_rg/runtime_proofs/docx_manifest/docx_manifest.json",
        output_dir=out,
        output_docx=out / _EXPECTED_DOCX_BASENAME,
    )


def build_docx_from_final_resume(paths: DocxRendererPaths | None = None) -> dict[str, Any]:
    if Document is None:
        raise RuntimeError(
            "python-docx is required. Install workspace deps (pip install python-docx) for DOCX emission.",
        )
    paths = paths or resolve_docx_renderer_paths()
    repo = paths.repo_root
    utc_now = datetime.now(timezone.utc).isoformat()

    fb = json.loads(paths.final_resume_json.read_text(encoding="utf-8"))
    mb = json.loads(paths.docx_manifest_json.read_text(encoding="utf-8"))

    by_section_id = {
        str(s["section_id"]): s for s in (fb.get("sections") or []) if isinstance(s, dict) and "section_id" in s
    }

    dm_order = mb.get("section_render_order") or []
    narrative_order = [str(x) for x in dm_order]
    profiles_raw = mb.get("section_profiles") or []
    prof_by_id = {}
    if isinstance(profiles_raw, list):
        prof_by_id = {
            str(p.get("section_id")): p
            for p in profiles_raw
            if isinstance(p, dict) and p.get("section_id")
        }

    invariant_order_raw = mb.get("locked_copy_invariants_projection_order")
    if isinstance(invariant_order_raw, list):
        invariant_order = [str(x) for x in invariant_order_raw]
    else:
        invariant_order = ["company_names", "titles", "locations", "dates"]

    fr_bytes_sha256 = _sha256_file(paths.final_resume_json.resolve())
    manifest_bytes_sha256 = _sha256_file(paths.docx_manifest_json.resolve())

    render_evidence: list[dict[str, Any]] = []

    paths.output_dir.mkdir(parents=True, exist_ok=True)
    doc = Document()

    ident = fb.get("candidate_identity") if isinstance(fb.get("candidate_identity"), dict) else {}
    cname = str(ident.get("candidate_name") or "").strip()
    hc_raw = ident.get("header_contact") if isinstance(ident.get("header_contact"), dict) else {}
    order_ct = ("phone", "email", "linkedin", "github", "location")
    contact_bits = [str(hc_raw[k]).strip() for k in order_ct if hc_raw.get(k) and str(hc_raw[k]).strip()]
    contact_line = " | ".join(contact_bits) if contact_bits else ""
    prem_blocks: list[str] = []
    if cname:
        prem_blocks.append(cname)
    if contact_line:
        prem_blocks.append(contact_line)
    if cname:
        _paragraph_safe(doc, cname, style_name=_heading_style_name(1))
    if contact_line:
        _paragraph_safe(doc, contact_line, style_name="Normal")
    if cname or contact_line:
        doc.add_paragraph()
        render_evidence.append(
            {
                "section_id": "candidate_identity",
                "section_kind": "base_resume_verbatim",
                "outline_level_emitted_for_section_heading": 1,
                "block_count": len(prem_blocks),
                "expected_plaintext_blocks": prem_blocks,
            },
        )

    for sid in narrative_order:
        section_blob = by_section_id.get(sid)
        if section_blob is None:
            raise ValueError(f"final_resume missing canonical section id {sid}")
        prof = prof_by_id.get(sid) or {}
        sm_raw = prof.get("style_mapping")
        style_map = sm_raw if isinstance(sm_raw, dict) else {}

        lvl = style_map.get("section_heading_outline_level")
        lvl_i = int(lvl) if isinstance(lvl, int) else 2
        title = str(prof.get("human_section_title_hint") or sid)

        bullet_ext = (
            style_map.get("bullet_list_style_external_id")
            if isinstance(style_map, dict)
            else None
        )

        _paragraph_safe(doc, title, style_name=_heading_style_name(lvl_i))

        plain_blocks = extract_plaintext_blocks(sid=sid, section_blob=section_blob)

        render_evidence.append(
            {
                "section_id": sid,
                "outline_level_emitted_for_section_heading": lvl_i,
                "block_count": 0,
                "expected_plaintext_blocks": plain_blocks,
            },
        )
        evid_entry = render_evidence[-1]

        if sid == "competencies":
            _render_competency_section(doc, section_blob)
            evid_entry["block_count"] = len(plain_blocks)
        elif bullet_ext == "resume_standard_bullet" and sid in {"unify_bullets", "ibm_bullets"}:
            for ln in plain_blocks:
                _paragraph_safe(doc, ln, style_name="List Bullet")
                evid_entry["block_count"] += 1
        else:
            for ln in plain_blocks:
                _paragraph_safe(doc, ln, style_name="Normal")
                evid_entry["block_count"] += 1

        if evid_entry["block_count"] < 1:
            raise ValueError(f"Section {sid} emitted zero plaintext blocks.")

    lc_inv_src = fb.get("locked_copy_invariants")
    lc_inv_src = lc_inv_src if isinstance(lc_inv_src, dict) else {}

    inv_levels: list[int] = []
    for nid in narrative_order:
        pr = prof_by_id.get(nid) or {}
        smap = pr.get("style_mapping")
        if isinstance(smap, dict) and isinstance(smap.get("section_heading_outline_level"), int):
            inv_levels.append(int(smap["section_heading_outline_level"]))
    inv_heading_level = max(inv_levels) if inv_levels else 2

    for inv_id in invariant_order:
        inv_blob_raw = lc_inv_src.get(inv_id)
        inv_blob_section = inv_blob_raw if isinstance(inv_blob_raw, dict) else {}
        ct = inv_blob_section.get("copied_text_exact")
        if ct is None:
            raise ValueError(f"Missing invariant copied_text_exact for {inv_id}")
        hdr = f"Invariant ΓÇö {inv_id}"
        _paragraph_safe(doc, hdr, style_name=_heading_style_name(inv_heading_level))
        plain_blocks_inv = [str(ct)]
        _paragraph_safe(doc, str(ct), style_name="Normal")
        render_evidence.append(
            {
                "section_id": inv_id,
                "section_kind": "locked_copy_invariant",
                "outline_level_emitted_for_section_heading": inv_heading_level,
                "block_count": 1,
                "expected_plaintext_blocks": plain_blocks_inv,
            },
        )

    doc.save(str(paths.output_docx.resolve()))

    render_manifest_blob: dict[str, Any] = {
        "render_manifest_id": "docx_render_manifest_v1",
        "constructed_at_utc": utc_now,
        "constructor_module": "apps_rg.runtime.render.docx_renderer",
        "sources": {
            "final_resume_json": paths.rel(paths.final_resume_json),
            "docx_manifest_json": paths.rel(paths.docx_manifest_json),
            "final_resume_sha256_bytes": fr_bytes_sha256,
            "docx_manifest_sha256_bytes": manifest_bytes_sha256,
            "final_resume_hash_logical": str(fb.get("final_resume_hash") or ""),
        },
        "section_render_order": narrative_order,
        "locked_copy_invariants_render_order_emitted": invariant_order,
        "output_docx": paths.rel(paths.output_docx),
        "verification": {
            "provider_calls_made": False,
            "qwen_calls_made": False,
            "judge_calls_made": False,
            "semantic_rewrite_attempted": False,
        },
        "render_evidence": render_evidence,
    }

    receipt_path = paths.output_dir / "docx_render_receipt.json"
    x2_path = paths.output_dir / "docx_render_x2_gate_outputs.json"
    mf_path = paths.output_dir / "docx_render_manifest.json"

    mf_path.write_text(
        json.dumps(render_manifest_blob, ensure_ascii=False, indent=2, sort_keys=False) + "\n",
        encoding="utf-8",
    )
    receipt_path.write_text("{}\n", encoding="utf-8")

    gate_results = run_docx_render_x2_gates(
        repo_root=repo,
        render_manifest_blob=render_manifest_blob,
        final_resume_blob=fb,
        docx_manifest_blob=mb,
        final_resume_path=paths.final_resume_json.resolve(),
        docx_manifest_path=paths.docx_manifest_json.resolve(),
        docx_output_path=paths.output_docx.resolve(),
        receipt_path=receipt_path.resolve(),
        output_dir=paths.output_dir.resolve(),
        expected_docx_basename=_EXPECTED_DOCX_BASENAME,
    )
    all_pass = gates_all_pass(gate_results)
    failed_gate_ids = failures(gate_results)

    x2_path.write_text(
        json.dumps(
            gate_records_to_blob(
                gate_results,
                evaluated_at_utc=utc_now,
                all_pass_res=all_pass,
                failed_gate_ids_res=failed_gate_ids,
            ),
            ensure_ascii=False,
            indent=2,
            sort_keys=False,
        )
        + "\n",
        encoding="utf-8",
    )

    receipt_path.write_text(
        json.dumps(
            {
                "receipt_id": "docx_render_receipt_v1",
                "written_at_utc": utc_now,
                "output_docx_rel": paths.rel(paths.output_docx),
                "manifest_json_rel": paths.rel(mf_path),
                "x2_gate_outputs_json_rel": paths.rel(x2_path),
                "gates_all_pass": all_pass,
                "failed_gate_ids": failed_gate_ids,
            },
            ensure_ascii=False,
            indent=2,
            sort_keys=False,
        )
        + "\n",
        encoding="utf-8",
    )

    return {
        "docx_path": paths.output_docx,
        "manifest_path": mf_path,
        "x2_path": x2_path,
        "receipt_path": receipt_path,
        "gates_all_pass": all_pass,
        "failed_gate_ids": failed_gate_ids,
        "final_resume_hash_logical": fb.get("final_resume_hash"),
        "manifest_bytes_sha256": manifest_bytes_sha256,
    }

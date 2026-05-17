"""Template-free JSON → DOCX export for apps_rg (post-L2 resume payload).

Handles common shapes: legacy ops_scripts JSON, rg_output_schema-ish ``sections``,
and arbitrary dicts (JSON snapshot fallback).

Body section order (after name / contact / optional tailoring line) matches canonical
base resume / merged ``rg_output.sections`` key order: executive summary, professional
experience, engineering & platform competencies, education, certifications & credentials.
"""
from __future__ import annotations

import json
from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Pt

from apps_rg.runtime.section_display_labels import (
    CERTIFICATIONS_AND_CREDENTIALS_HEADING,
    ENGINEERING_PLATFORM_COMPETENCIES_HEADING,
)


def _is_provider_stub_payload(data: dict[str, Any]) -> bool:
    """True for LLM stub envelopes and other non-resume dicts we should not json-line-dump."""
    if data.get("stub_response") is True:
        return True
    resume_hint_keys = (
        "sections",
        "experience",
        "contact_info",
        "candidate_name",
        "name",
        "headline",
        "headline_line",
        "tagline",
        "summary",
        "executive_summary",
        "skills",
        "education",
    )
    if any(k in data for k in resume_hint_keys):
        return False
    # Small opaque dicts (e.g. {hash, note, error}) — treat as metadata-only
    if len(data) <= 6 and not any(isinstance(v, (list, dict)) for v in data.values()):
        return True
    return False


def _render_stub_or_metadata_body(
    doc: Document,
    data: dict[str, Any],
    *,
    target_role: str,
    target_company: str,
    title_line: str,
) -> None:
    """Readable export when there is no structured resume (stub provider, errors, receipts)."""
    _add_heading(doc, title_line)
    if target_company.strip():
        _add_para(doc, f"Target organization: {target_company.strip()}")
    if target_role.strip():
        _add_para(doc, f"Target role: {target_role.strip()}")
    doc.add_paragraph()

    _add_h2(doc, "GENERATION STATUS")
    if data.get("stub_response") is True:
        _add_para(
            doc,
            "This run completed with a provider stub. No LLM-authored résumé sections "
            "were produced; attach a live model profile to generate full content.",
        )
    note = data.get("note")
    if note:
        _add_para(doc, str(note))
    err = data.get("error")
    if err:
        _add_para(doc, f"Error: {err}", bold=True)
    h = data.get("hash")
    if h:
        _add_para(doc, f"Content fingerprint: {h}")

    # Any remaining primitive keys (excluding those already shown)
    shown = {"stub_response", "note", "error", "hash"}
    extras = [
        (k, v)
        for k, v in sorted(data.items())
        if k not in shown and not isinstance(v, (dict, list))
    ]
    if extras:
        doc.add_paragraph()
        _add_h2(doc, "ADDITIONAL FIELDS")
        for k, v in extras:
            _add_para(doc, f"{k}: {v}")


def _add_heading(doc: Document, text: str) -> None:
    p = doc.add_heading(text, level=1)
    for r in p.runs:
        r.font.size = Pt(14)


def _add_h2(doc: Document, text: str) -> None:
    p = doc.add_heading(text, level=2)
    for r in p.runs:
        r.font.size = Pt(11)


def _add_para(doc: Document, text: str, *, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(10)


def _format_contact_line(contact: dict[str, Any] | None) -> str:
    """Single header line: phone | email | linkedin | github | location."""
    if not isinstance(contact, dict) or not contact:
        return ""
    order = ("phone", "email", "linkedin", "github", "location")
    bits: list[str] = []
    for k in order:
        v = contact.get(k)
        if v is None:
            continue
        s = str(v).strip()
        if s:
            bits.append(s)
    return " | ".join(bits)


def _format_education_display_line(entry: dict[str, Any]) -> str:
    """One résumé-style clause: ``degree, institution[, year] [(honors)]``."""
    degree = str(entry.get("degree", "") or "").strip()
    inst = str(entry.get("institution", "") or "").strip()
    year_raw = entry.get("year", entry.get("graduation_year", ""))
    year = str(year_raw).strip() if year_raw not in (None, "") else ""
    honors = str(entry.get("honors", "") or "").strip()
    parts: list[str] = [p for p in (degree, inst) if p]
    if year:
        parts.append(year)
    line = ", ".join(parts)
    if honors:
        line = f"{line} ({honors})"
    return line


def _format_certification_display_line(entry: dict[str, Any]) -> str:
    """``name[, issuer][, date/year]`` — matches résumé comma style (issuer omitted when absent)."""
    name = str(entry.get("name", "") or "").strip()
    issuer = str(entry.get("issuer", "") or "").strip()
    yr_raw = entry.get("date", entry.get("year", ""))
    yr = str(yr_raw).strip() if yr_raw not in (None, "") else ""
    parts: list[str] = [name] if name else []
    if issuer:
        parts.append(issuer)
    if yr:
        parts.append(yr)
    return ", ".join(parts)


def _candidate_display_name(data: dict[str, Any], fallback_name: str) -> str:
    """Resolve résumé heading name without rewriting (prefers candidate_name / header.name)."""
    if isinstance(data.get("candidate_name"), str) and data["candidate_name"].strip():
        return data["candidate_name"].strip()
    if isinstance(data.get("name"), str) and data["name"].strip():
        return data["name"].strip()
    hdr = data.get("header") if isinstance(data.get("header"), dict) else {}
    hn = str(hdr.get("name") or "").strip()
    if hn:
        return hn
    ci = data.get("contact_info") if isinstance(data.get("contact_info"), dict) else {}
    cn = str(ci.get("name") or "").strip()
    if cn:
        return cn
    fb = (fallback_name or "").strip()
    return fb if fb else "Candidate"


def _normalize_for_render(
    data: dict[str, Any],
    *,
    fallback_name: str,
    fallback_role: str,
    fallback_company: str,
) -> dict[str, Any]:
    if "sections" in data and isinstance(data["sections"], dict):
        sec = data["sections"]
        summary = sec.get("summary") or {}
        summary_text = ""
        if isinstance(summary, dict):
            summary_text = str(summary.get("text", "") or "")
        exp = sec.get("experience") or []
        edu = sec.get("education") or []
        headline_line = str(data.get("headline_line") or "").strip()
        t_role = str(data.get("target_role") or fallback_role or "").strip()
        t_co = str(data.get("target_company") or fallback_company or "").strip()
        ci_payload = data.get("contact_info") if isinstance(data.get("contact_info"), dict) else {}
        hdr_top = data.get("header") if isinstance(data.get("header"), dict) else {}
        # Header (often canonical base slice) overwrites contact_info for the same keys — matches enrich/export parity.
        merged_contact: dict[str, Any] = {**ci_payload, **hdr_top}
        contact_line = _format_contact_line(merged_contact)
        targ_parts = [p for p in (t_role, t_co) if p]
        targeting_line = " — ".join(targ_parts) if targ_parts else ""
        tagline = headline_line or t_role

        raw_skills = sec.get("skills")
        skills_for_render: list[str] = []
        skills_categories: list[dict[str, Any]] = []
        if isinstance(raw_skills, list):
            for s in raw_skills:
                if isinstance(s, dict):
                    skills_for_render.append(str(s.get("label") or s.get("name") or s))
                else:
                    skills_for_render.append(str(s))
        elif isinstance(raw_skills, dict):
            for cat in raw_skills.get("categories") or []:
                if not isinstance(cat, dict):
                    continue
                nm = str(cat.get("name") or "").strip()
                terms = cat.get("items") or []
                if not isinstance(terms, list):
                    continue
                items_clean: list[str] = []
                for t in terms:
                    if isinstance(t, dict):
                        bit = str(t.get("text") or t.get("name") or t.get("label") or "").strip()
                    else:
                        bit = str(t).strip()
                    if bit and bit not in items_clean:
                        items_clean.append(bit)
                if not items_clean:
                    continue
                nm_key = nm or "Capabilities"
                skills_categories.append({"name": nm_key, "items": items_clean})
                skills_for_render.append(f"{nm_key}: {', '.join(items_clean)}")

        certs = sec.get("certifications") or []
        certs_list = certs if isinstance(certs, list) else []

        return {
            "name": _candidate_display_name(data, fallback_name),
            "tagline": tagline,
            "contact_line": contact_line,
            "targeting_line": targeting_line,
            "executive_summary": summary_text,
            "experience": exp if isinstance(exp, list) else [],
            "skills": skills_for_render,
            "skills_categories": skills_categories,
            "education": edu if isinstance(edu, list) else [],
            "certifications": certs_list,
            "raw_fallback": False,
        }

    ci = data.get("contact_info") if isinstance(data.get("contact_info"), dict) else {}
    hdr_top = data.get("header") if isinstance(data.get("header"), dict) else {}
    merged_legacy = {**ci, **hdr_top}
    name = _candidate_display_name(data, fallback_name)
    headline = str(data.get("headline") or data.get("tagline") or fallback_role or "")
    summary = str(data.get("summary") or data.get("executive_summary") or "")
    if isinstance(data.get("executive_summary"), dict):
        summary = str(data["executive_summary"].get("content", summary))
    contact = _format_contact_line(merged_legacy)
    t_role = str(data.get("target_role") or fallback_role or "").strip()
    t_co = str(data.get("target_company") or fallback_company or "").strip()
    targ_parts = [p for p in (t_role, t_co) if p]
    targeting_line = " — ".join(targ_parts) if targ_parts else ""
    certs_top = data.get("certifications") if isinstance(data.get("certifications"), list) else []
    return {
        "name": name,
        "tagline": headline,
        "contact_line": contact,
        "targeting_line": targeting_line,
        "executive_summary": summary,
        "experience": data.get("experience") or [],
        "skills": data.get("skills") or [],
        "skills_categories": [],
        "education": data.get("education") or [],
        "certifications": certs_top,
        "raw_fallback": False,
    }


def render_resume_dict_to_docx(
    data: dict[str, Any],
    output_path: Path,
    *,
    target_role: str = "",
    target_company: str = "",
) -> None:
    """Write ``data`` to ``output_path`` as a .docx (creates parent dirs)."""
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    if _is_provider_stub_payload(data):
        doc = Document()
        if target_company.strip():
            title_line = f"Résumé export — {target_company.strip()}"
        elif target_role.strip():
            title_line = f"Résumé export — {target_role.strip()}"
        else:
            title_line = "Résumé export (stub)"
        _render_stub_or_metadata_body(
            doc,
            data,
            target_role=target_role,
            target_company=target_company,
            title_line=title_line,
        )
        doc.save(output_path)
        return

    norm = _normalize_for_render(
        data,
        fallback_name=target_role and f"Candidate — {target_role}" or "Resume export",
        fallback_role=target_role,
        fallback_company=target_company,
    )
    if not any(
        [
            norm.get("executive_summary"),
            norm.get("experience"),
            norm.get("skills"),
            norm.get("skills_categories"),
            norm.get("education"),
            norm.get("tagline"),
            norm.get("certifications"),
        ]
    ):
        norm["raw_fallback"] = True

    doc = Document()

    _add_heading(doc, str(norm["name"]))
    if norm.get("tagline"):
        _add_para(doc, str(norm["tagline"]))
    if norm.get("contact_line"):
        _add_para(doc, str(norm["contact_line"]))
    if norm.get("targeting_line"):
        _add_para(doc, f"Tailored for: {norm['targeting_line']}")
    doc.add_paragraph()

    if norm.get("raw_fallback"):
        _add_h2(doc, "GENERATED CONTENT (JSON)")
        for line in json.dumps(data, ensure_ascii=False, indent=2).splitlines():
            _add_para(doc, line)
        doc.save(output_path)
        return

    if norm.get("executive_summary"):
        _add_h2(doc, "EXECUTIVE SUMMARY")
        _add_para(doc, str(norm["executive_summary"]))
        doc.add_paragraph()

    exp = norm.get("experience") or []
    if isinstance(exp, list) and exp:
        _add_h2(doc, "PROFESSIONAL EXPERIENCE")
        for role in exp:
            if not isinstance(role, dict):
                _add_para(doc, str(role))
                continue
            company = str(role.get("company", ""))
            title = str(role.get("title") or role.get("role", ""))
            dates = str(role.get("dates", ""))
            loc = str(role.get("location", ""))
            header = " | ".join(p for p in [company, loc] if p)
            if header:
                _add_para(doc, header, bold=True)
            line = " | ".join(p for p in [title, dates] if p)
            if line:
                _add_para(doc, line)
            narr = str(role.get("role_narrative") or "").strip()
            if narr:
                _add_para(doc, narr)
            bullets = role.get("bullets") or role.get("achievements") or []
            if isinstance(bullets, list):
                for b in bullets:
                    if isinstance(b, dict):
                        txt = str(b.get("description") or b.get("text") or "")
                    else:
                        txt = str(b)
                    if txt:
                        p = doc.add_paragraph(style="List Bullet")
                        r = p.add_run(txt)
                        r.font.size = Pt(10)
            doc.add_paragraph()

    # Order mirrors ``facts`` in base resume SSOT: employment → skills → education → certifications
    skills_cats = norm.get("skills_categories") or []
    skills_flat = norm.get("skills") or []
    if isinstance(skills_cats, list) and skills_cats:
        _add_h2(doc, ENGINEERING_PLATFORM_COMPETENCIES_HEADING)
        for cat in skills_cats:
            if not isinstance(cat, dict):
                continue
            nm = str(cat.get("name") or "").strip()
            items = cat.get("items") or []
            if nm:
                _add_para(doc, nm, bold=True)
            if isinstance(items, list):
                for it in items:
                    s = str(it).strip()
                    if not s:
                        continue
                    p = doc.add_paragraph(style="List Bullet")
                    r = p.add_run(s)
                    r.font.size = Pt(10)
        doc.add_paragraph()
    elif isinstance(skills_flat, list) and skills_flat:
        _add_h2(doc, ENGINEERING_PLATFORM_COMPETENCIES_HEADING)
        parts: list[str] = []
        for s in skills_flat:
            if isinstance(s, dict):
                parts.append(str(s.get("label") or s.get("name") or s))
            else:
                parts.append(str(s))
        _add_para(doc, " · ".join(parts))

    edu = norm.get("education") or []
    if isinstance(edu, list) and edu:
        _add_h2(doc, "EDUCATION")
        clauses: list[str] = []
        for e in edu:
            if isinstance(e, dict):
                disp = _format_education_display_line(e)
                clauses.append(disp if disp else json.dumps(e, ensure_ascii=False))
            else:
                clauses.append(str(e))
        _add_para(doc, " ".join(clauses))

    certs = norm.get("certifications") or []
    if isinstance(certs, list) and certs:
        _add_h2(doc, CERTIFICATIONS_AND_CREDENTIALS_HEADING)
        for c in certs:
            if isinstance(c, dict):
                disp = _format_certification_display_line(c)
                _add_para(doc, disp if disp else json.dumps(c, ensure_ascii=False))
            else:
                _add_para(doc, str(c))
        doc.add_paragraph()

    doc.save(output_path)


__all__ = ["render_resume_dict_to_docx"]

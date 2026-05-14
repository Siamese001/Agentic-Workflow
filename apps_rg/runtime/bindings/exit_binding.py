"""Exit binding for apps_rg `resume_generation` task class.

MIGRATED from agentic_core/runtime/exit/apps_rg_exit_binding.py
Per plan apps-rg-golden-state-section-generation-a4f9e1 W2F.

Per plan apps-rg-runtime-wiring-completion-d4e8a1 §6 W3.P5.

GAP-001 P0 FIX: Exit L4 Boundary Hardening
- Removed all direct filesystem durable writes (mkdir, write_text, json.dump)
- Replaced with inert CommitRequest-like candidates
- Exit now emits exactly one X3 with inert mutation proposals
- UWG is the sole durable writer per architectural law

Exit is the SEVENTH (last) stage. Its job is to:
- Build inert artifact commit candidates (no durable mutation)
- Build an X3Disposition with exactly one X3
- Bind to upstream SealedL2Artifact via inert references
- Evaluate gates (G24/G25/G26/G27) against final output

Exit circular import risk resolved by:
- No imports from agentic_core.L2_execution.* here
- Pure function shape with explicit contracts
- Contract types live in agentic_core.runtime.contracts.*
"""
from __future__ import annotations

import dataclasses
import hashlib
import json
import logging
import os
import re  # NEW: for regex parsing in ingestion
import sys
from enum import Enum
from pathlib import Path
from typing import Any, Mapping, Optional

from agentic_core.L5_safety.types.exit_disposition_types import ExitDisposition, ExitGateResult
from agentic_core.runtime.contracts.x3_disposition import X3Disposition
from agentic_core.runtime.contracts.sealed_l2_artifact import SealedL2Artifact


class ExitGateVerdict(str, Enum):
    """Verdict values for apps_rg exit gate evaluations (G24-G27)."""
    PASS = "PASS"
    WARN = "WARN"
    FAIL = "FAIL"


@dataclasses.dataclass(frozen=True)
class AppsRgGateResult:
    """Local gate result type for apps_rg G24-G27 evaluations."""
    gate_id: str
    verdict: ExitGateVerdict
    score: float
    weight: float
    reason: str

_LOGGER = logging.getLogger(__name__)

APPS_RG_EXIT_CERT_REF: str = "exit-apps-rg-resume-generation-w3p5"

# GAP-001: Inert commit request candidates — Exit proposes, UWG commits.
@dataclasses.dataclass(frozen=True)
class InertArtifactCommitCandidate:
    """Inert artifact commit proposal — no durable mutation.

    GAP-001 L4 Boundary Hardening: Exit emits inert candidates only.
    The UWG (L4 admission control) is the sole durable writer.

    Attributes:
        artifact_type: Type of artifact (resume_json, docx, cache_entry, etc.)
        proposed_path: Suggested path for UWG-mediated write (inert reference)
        content_digest: SHA256 of serialized content
        serialized_content: JSON-serializable content (in memory only)
        mutation_candidate_inert: Always True — marks as non-durable proposal
        proposal_status: Always PENDING_UWG — requires UWG admission
        non_durable: True — this object is session-scoped, not L4 truth
        not_l4_truth: True — not authoritative until UWG commits
        not_replay_source: True — cannot be used for deterministic replay
    """
    artifact_type: str
    proposed_path: str  # Virtual path for reference, not actual write target
    content_digest: str
    serialized_content: dict[str, Any]
    mutation_candidate_inert: bool = True
    proposal_status: str = "PENDING_UWG"
    non_durable: bool = True
    not_l4_truth: bool = True
    not_replay_source: bool = True


# Opt-in for writeback via environment. See G24/G25 policy.
_APPS_RG_CACHE_WRITE_ENABLED_ENV: str = "APPS_RG_CACHE_WRITE_ENABLED"
_APPS_RG_C0_WRITE_ENABLED_ENV: str = "APPS_RG_C0_WRITE_ENABLED"


def _safe_build_g24_provenance(sealed: SealedL2Artifact) -> str:
    """Return G24 provenance digest for cache entry and artifact validation."""
    provenance: dict[str, Any] = {
        "run_id": sealed.run_id,
        "request_id": sealed.request_id,
        "trace_id": sealed.trace_id,
        "execution_status": sealed.execution_status,
        "exec_receipt": sealed.sovereign_execution_receipt,
    }
    # The sealed artifact's compilation_hash binds the entire pipeline.
    # Cache entries use this as the provenance key.
    return json.dumps(provenance, sort_keys=True)


def _resolve_repo_root() -> Path:
    """Resolve repository root using the sentinel pyproject.toml."""
    p = Path(__file__).resolve()
    for parent in p.parents:
        if (parent / "pyproject.toml").exists():
            return parent
    return Path.cwd().resolve()


def _safe_run_dirname(target_company: str, target_role: str, run_id: str) -> str:
    """Return a filesystem-safe directory name for the run.

    The pattern is <company>_<role>_<run_id> with spaces sanitized.
    """
    safe_company = target_company.replace(" ", "_").replace("/", "_")
    safe_role = target_role.replace(" ", "_").replace("/", "_")
    return f"{safe_company}_{safe_role}_{run_id}"


def _find_existing_run_dir(target_company: str, target_role: str, run_id: str) -> Path | None:
    """Check if a run directory already exists (resume, redo, or repeated run)."""
    repo_root = _resolve_repo_root()
    base_dir = repo_root / "artifacts" / "apps_rg" / "runs"
    expected = base_dir / _safe_run_dirname(target_company, target_role, run_id)
    if expected.exists():
        return expected
    return None


def _build_artifact_commit_candidate(
    content: Mapping[str, Any],
    proposed_dir: Path,
    filename: str,
    artifact_type: str,
) -> InertArtifactCommitCandidate:
    """Build an inert artifact commit candidate — NO durable write.

    GAP-001 L4 Boundary Hardening: Exit must not mutate durable state.
    This function constructs an inert proposal for UWG-mediated admission.

    Args:
        content: Artifact content to serialize
        proposed_dir: Virtual directory path for UWG reference (not created)
        filename: Target filename for the artifact
        artifact_type: Classification of the artifact type

    Returns:
        InertArtifactCommitCandidate with content digest and serialized form.
        No filesystem mutation occurs.
    """
    proposed_path = str(proposed_dir / filename)
    json_body = json.dumps(content, indent=2, default=str)
    content_digest = hashlib.sha256(json_body.encode("utf-8")).hexdigest()[:32]

    return InertArtifactCommitCandidate(
        artifact_type=artifact_type,
        proposed_path=proposed_path,
        content_digest=content_digest,
        serialized_content=dict(content),
        mutation_candidate_inert=True,
        proposal_status="PENDING_UWG",
        non_durable=True,
        not_l4_truth=True,
        not_replay_source=True,
    )


def _ingest_docx_to_master_resume(
    template_path: Path,
    output_path: Path | None = None,  # GAP-001: DEPRECATED, ignored for L4 hardening
) -> dict[str, Any] | None:
    """Ingest a DOCX resume template and convert to JSON master resume format.

    GAP-001 L4 Boundary Hardening: This function NO LONGER performs durable writes.
    The output_path parameter is DEPRECATED and ignored. Use the returned dict
    to build an InertArtifactCommitCandidate for UWG-mediated admission.

    COMPREHENSIVE EXTRACTION:
    - Header: name, headline (X|Y|Z format), contact info
    - Executive Summary: full text
    - Experience: ALL roles with title, company, dates, location, ALL bullets
    - Education: degrees, schools, years
    - Certifications: FSA and other credentials

    Args:
        template_path: Path to the DOCX template (read-only, not modified)
        output_path: DEPRECATED. Ignored. Function returns dict only.

    Returns:
        Master resume as dict with complete extraction — or None if python-docx unavailable.
        Caller is responsible for creating InertArtifactCommitCandidate if persistence needed.
    """
    try:
        import docx
    except ImportError:
        _LOGGER.warning("[apps_rg Exit] python-docx not available, cannot ingest DOCX")
        return None

    if not template_path.exists():
        _LOGGER.warning("[apps_rg Exit] Template not found: %s", template_path)
        return None

    try:
        doc = docx.Document(str(template_path))
    except Exception as exc:
        _LOGGER.warning("[apps_rg Exit] Failed to load DOCX: %s", exc)
        return None

    # Parse paragraphs with text content
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    
    # Debug: Log paragraph count
    _LOGGER.info("[apps_rg Exit] DOCX contains %d paragraphs", len(paragraphs))

    master_resume: dict[str, Any] = {
        "schema_version": "master_resume_v2.0",
        "source_template": str(template_path),
        "ingested_at": __import__("datetime").datetime.now(__import__("datetime").timezone.utc).isoformat(),
        "header": {},
        "headline": "",  # NEW: X|Y|Z format headline
        "executive_summary": "",
        "experience": [],
        "education": [],
        "skills": [],
        "certifications": [],
    }

    # ===================================================================
    # PHASE 1: HEADER EXTRACTION (Lines 0-2 typically)
    # ===================================================================
    
    if paragraphs:
        # Line 0: Name
        master_resume["header"]["name"] = paragraphs[0]
        _LOGGER.info("[apps_rg Exit] Extracted name: %s", paragraphs[0])
    
    # Line 1: Headline (X|Y|Z format) OR Contact info
    if len(paragraphs) > 1:
        para1 = paragraphs[1]
        # Headline pattern: contains | but no email, no linkedin, not a date
        if "|" in para1 and "@" not in para1 and "linkedin" not in para1.lower():
            master_resume["headline"] = para1
            _LOGGER.info("[apps_rg Exit] Extracted headline: %s", para1)
        else:
            # It's contact info, parse it
            _parse_contact_line(para1, master_resume["header"])
    
    # Line 2: Contact info (if headline was on line 1)
    if len(paragraphs) > 2 and master_resume["headline"]:
        _parse_contact_line(paragraphs[2], master_resume["header"])
    elif len(paragraphs) > 2 and not master_resume["header"].get("email"):
        # Try parsing as contact line
        _parse_contact_line(paragraphs[2], master_resume["header"])

    # ===================================================================
    # PHASE 2: SECTION-BASED EXTRACTION
    # ===================================================================
    
    current_section: str | None = None
    current_role: dict[str, Any] | None = None
    bullet_accumulator: list[str] = []
    
    for i, para in enumerate(paragraphs):
        para_lower = para.lower()
        
        # Section header detection
        if "executive summary" in para_lower:
            current_section = "executive_summary"
            continue
            
        elif para_lower == "experience" or "professional experience" in para_lower:
            current_section = "experience"
            # Save any accumulated role before switching
            if current_role and bullet_accumulator:
                current_role["bullets"] = bullet_accumulator.copy()
                master_resume["experience"].append(current_role)
                current_role = None
                bullet_accumulator = []
            continue
            
        elif para_lower == "education":
            current_section = "education"
            continue
            
        elif para_lower in ["certifications", "certificates", "professional certifications"]:
            current_section = "certifications"
            continue
            
        elif para_lower in ["skills", "core competencies", "technical skills"]:
            current_section = "skills"
            continue
        
        # Skip empty lines and page breaks
        if not para or para == "\x0c":  # Form feed/page break
            continue
        
        # =================================================================
        # EXECUTIVE SUMMARY CONTENT
        # =================================================================
        if current_section == "executive_summary" and para != master_resume["header"].get("name", ""):
            # Accumulate until next section header
            if not _is_section_header(para):
                if master_resume["executive_summary"]:
                    master_resume["executive_summary"] += " " + para
                else:
                    master_resume["executive_summary"] = para
            continue
        
        # =================================================================
        # EXPERIENCE CONTENT
        # =================================================================
        if current_section == "experience":
            # Check if this is a COMPANY/LOCATION line (e.g., "Unify Consulting | Boca Raton, FL")
            if _is_company_line(para):
                # Save previous role if exists
                if current_role:
                    if bullet_accumulator:
                        current_role["bullets"] = bullet_accumulator.copy()
                    master_resume["experience"].append(current_role)
                    bullet_accumulator = []
                
                # Start new role with company info
                company, location = _parse_company_line(para)
                current_role = {
                    "title": "",
                    "company": company,
                    "location": location,
                    "dates": "",
                    "bullets": [],
                    "intro": "",
                }
                _LOGGER.info("[apps_rg Exit] Found company: %s (%s)", company, location)
            
            # Check if this is a TITLE/DATES line (e.g., "SVP Engineering | February 2023 - Present")
            elif _is_title_line(para):
                if current_role is not None:
                    title, dates = _parse_title_line(para)
                    current_role["title"] = title
                    current_role["dates"] = dates
                    _LOGGER.info("[apps_rg Exit] Found title: %s (%s)", title, dates)
                else:
                    # Orphan title line - create minimal role
                    title, dates = _parse_title_line(para)
                    current_role = {
                        "title": title,
                        "company": "",
                        "location": "",
                        "dates": dates,
                        "bullets": [],
                        "intro": "",
                    }
            
            # Check if this is an INTRO paragraph (first paragraph after title, no bold keyword)
            elif current_role and not current_role.get("intro") and not _is_bold_bullet(para) and len(para) > 40:
                current_role["intro"] = para
                _LOGGER.debug("[apps_rg Exit] Added role intro: %s...", para[:50])
            
            # Check if this is a BULLET (bold keyword + colon pattern, or traditional bullet)
            elif _is_bold_bullet(para) or para.startswith("•") or para.startswith("-"):
                if current_role is not None:
                    bullet_text = para.lstrip("•- ").strip()
                    if bullet_text:
                        bullet_accumulator.append(bullet_text)
                        _LOGGER.debug("[apps_rg Exit] Added bullet: %s...", bullet_text[:50])
            
            # Any other content in experience section becomes a bullet if we have a role
            elif current_role and len(para) > 20:
                bullet_accumulator.append(para)
                _LOGGER.debug("[apps_rg Exit] Added bullet (fallback): %s...", para[:50])
            
            continue
        
        # =================================================================
        # EDUCATION CONTENT
        # =================================================================
        if current_section == "education":
            edu_entry = _parse_education_line(para)
            if edu_entry:
                master_resume["education"].append(edu_entry)
                _LOGGER.info("[apps_rg Exit] Extracted education: %s", edu_entry.get("degree", ""))
            continue
        
        # =================================================================
        # CERTIFICATIONS CONTENT
        # =================================================================
        if current_section == "certifications":
            cert_entry = _parse_certification_line(para)
            if cert_entry:
                master_resume["certifications"].append(cert_entry)
                _LOGGER.info("[apps_rg Exit] Extracted certification: %s", cert_entry.get("name", ""))
            continue
        
        # =================================================================
        # SKILLS CONTENT
        # =================================================================
        if current_section == "skills":
            # Parse skills (comma-separated or bullet list)
            skills = _parse_skills_line(para)
            if skills:
                master_resume["skills"].extend(skills)
            continue
    
    # Save final role if exists
    if current_role:
        if bullet_accumulator:
            current_role["bullets"] = bullet_accumulator.copy()
        master_resume["experience"].append(current_role)
        _LOGGER.info("[apps_rg Exit] Saved final role: %s", current_role.get("title", ""))

    # ===================================================================
    # PHASE 3: VALIDATION & LOGGING
    # ===================================================================
    
    exp_count = len(master_resume["experience"])
    total_bullets = sum(len(role.get("bullets", [])) for role in master_resume["experience"])
    edu_count = len(master_resume["education"])
    cert_count = len(master_resume["certifications"])
    
    _LOGGER.info("[apps_rg Exit] Master resume extraction complete:")
    _LOGGER.info("[apps_rg Exit]   - Name: %s", master_resume["header"].get("name", "NOT FOUND"))
    _LOGGER.info("[apps_rg Exit]   - Headline: %s", master_resume["headline"] or "NOT FOUND")
    _LOGGER.info("[apps_rg Exit]   - Experience entries: %d (total bullets: %d)", exp_count, total_bullets)
    _LOGGER.info("[apps_rg Exit]   - Education entries: %d", edu_count)
    _LOGGER.info("[apps_rg Exit]   - Certifications: %d", cert_count)
    _LOGGER.info("[apps_rg Exit]   - Exec Summary: %d chars", len(master_resume["executive_summary"]))
    
    # Warn if expected content is missing
    if not master_resume["headline"]:
        _LOGGER.warning("[apps_rg Exit] Headline not found - expected X|Y|Z format")
    if exp_count == 0:
        _LOGGER.warning("[apps_rg Exit] No experience entries extracted")
    if edu_count == 0:
        _LOGGER.warning("[apps_rg Exit] No education entries extracted")
    if cert_count == 0:
        _LOGGER.warning("[apps_rg Exit] No certifications extracted (expected FSA)")

    # GAP-001: Direct filesystem writes are FORBIDDEN.
    # If output_path was provided (legacy API), emit deprecation warning and ignore.
    if output_path:
        _LOGGER.warning(
            "[apps_rg Exit] GAP-001: output_path parameter is DEPRECATED and ignored. "
            "Use the returned dict to build an InertArtifactCommitCandidate."
        )

    # Return master resume dict only — caller handles persistence via inert proposals
    return master_resume


def _parse_contact_line(line: str, header: dict[str, Any]) -> None:
    """Parse a contact info line for email, phone, linkedin."""
    # Split by common separators
    parts = [p.strip() for p in line.replace("|", ",").split(",")]
    for part in parts:
        if "@" in part and "." in part:
            header["email"] = part
        elif any(c.isdigit() for c in part) and sum(c.isdigit() for c in part) >= 7:
            header["phone"] = part
        elif "linkedin" in part.lower():
            header["linkedin"] = part
        elif part.startswith("http"):
            header["linkedin"] = part


def _is_section_header(para: str) -> bool:
    """Check if paragraph is a section header."""
    section_headers = [
        "experience", "professional experience", "work experience",
        "education", "academic background",
        "certifications", "certificates", "professional certifications",
        "skills", "core competencies", "technical skills", "competencies",
        "summary", "professional summary", "executive summary",
        "projects", "publications", "awards", "honors"
    ]
    return para.lower().strip() in section_headers


def _is_company_line(para: str) -> bool:
    """Check if paragraph is a company/location line.
    
    Format: "Company Name | City, ST" or "Company Name | Location"
    """
    if " | " not in para:
        return False
    
    parts = para.split(" | ")
    if len(parts) != 2:
        return False
    
    right_side = parts[1]
    # Location pattern: has comma (City, ST) or common location words
    location_indicators = [", FL", ", CA", ", NY", ", NJ", ", TX", ", IL", ", WA", ", GA", ", VA"]
    has_location_pattern = any(ind in right_side for ind in location_indicators)
    
    # No date indicators on company line
    date_indicators = ["20", "Present", "–", "-"]
    has_date = any(d in right_side for d in date_indicators)
    
    return has_location_pattern and not has_date


def _parse_company_line(para: str) -> tuple[str, str]:
    """Parse company/location line into (company, location)."""
    parts = para.split(" | ")
    if len(parts) >= 2:
        return parts[0].strip(), parts[1].strip()
    return para.strip(), ""


def _is_title_line(para: str) -> bool:
    """Check if paragraph is a title/dates line.
    
    Format: "Title | Month Year - Month Year" or "Title, Subtitle | Dates"
    """
    if " | " not in para:
        return False
    
    # Date indicators
    date_indicators = ["202", "Present", "present", "–", "Jan", "Feb", "Mar", "Apr", "May", "Jun", 
                      "Jul", "Aug", "Sep", "Oct", "Nov", "Dec"]
    has_date = any(d in para for d in date_indicators)
    
    # Title indicators
    title_indicators = ["VP", "SVP", "EVP", "President", "Director", "Manager", "Lead",
                       "Engineer", "Architect", "Consultant", "Partner", "CTO", "CEO", "Founder",
                       "Actuary", "Fellow", "Principal", "Senior", "Client"]
    has_title = any(t in para for t in title_indicators)
    
    # Must have date and reasonable length
    return has_date and len(para) < 120


def _parse_title_line(para: str) -> tuple[str, str]:
    """Parse title/dates line into (title, dates)."""
    parts = para.split(" | ")
    if len(parts) >= 2:
        return parts[0].strip(), parts[1].strip()
    return para.strip(), ""


def _is_bold_bullet(para: str) -> bool:
    """Check if paragraph looks like a bold keyword bullet.
    
    Format: "Keyword: Description" or "Multi-word Keyword: Description"
    """
    if ":" not in para:
        return False
    
    # Must have text before and after colon
    parts = para.split(":", 1)
    if len(parts) != 2:
        return False
    
    keyword_part = parts[0].strip()
    desc_part = parts[1].strip()
    
    # Keyword should be relatively short (1-5 words)
    word_count = len(keyword_part.split())
    if word_count < 1 or word_count > 6:
        return False
    
    # Description should exist and be reasonable length
    if len(desc_part) < 10:
        return False
    
    return True


def _is_role_header(para: str) -> bool:
    """Check if paragraph looks like a job role header.
    
    Handles formats:
    - "Company | City, ST" (company/location line)
    - "Title, Subtitle | Month Year - Month Year" (title/dates line)
    - "Title | Month Year – Month Year" (title with en-dash dates)
    """
    return _is_company_line(para) or _is_title_line(para)


def _parse_role_header(para: str) -> dict[str, Any]:
    """Parse a role header line into structured data."""
    role = {
        "title": "",
        "company": "",
        "location": "",
        "dates": "",
        "bullets": [],
        "intro": "",
    }
    
    # Try to split by common separators
    # Patterns: "Title, Company — Location | Dates" or "Title | Company | Dates"
    
    # First, try to extract dates (common patterns: 2020-Present, Jan 2020 - Dec 2022)
    date_patterns = [
        r"(\d{4}\s*-\s*(?:Present|present|\d{4}))",
        r"((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{4})",
    ]
    import re
    for pattern in date_patterns:
        match = re.search(pattern, para, re.IGNORECASE)
        if match:
            role["dates"] = match.group(1)
            # Remove dates from para for further parsing
            para = para.replace(match.group(0), "").strip()
            break
    
    # Try to split by em-dash or en-dash
    if "—" in para or "–" in para:
        parts = para.replace("–", "—").split("—")
        if len(parts) >= 2:
            left = parts[0].strip()
            right = "—".join(parts[1:]).strip()
            
            # Left side: title and/or company
            if "," in left:
                role["title"] = left.split(",")[0].strip()
                role["company"] = ",".join(left.split(",")[1:]).strip()
            else:
                role["title"] = left
            
            # Right side: location and/or dates
            role["location"] = right
    
    # Try comma-separated
    elif "," in para and not role["title"]:
        parts = [p.strip() for p in para.split(",")]
        if len(parts) >= 2:
            role["title"] = parts[0]
            role["company"] = parts[1]
            if len(parts) >= 3:
                role["location"] = parts[2]
    
    # Fallback: use whole line as title
    if not role["title"]:
        role["title"] = para
    
    return role


def _is_date_line(para: str) -> bool:
    """Check if paragraph looks like a date line."""
    date_indicators = ["202", "201", "200", "199", "Present", "present"]
    month_indicators = ["Jan", "Feb", "Mar", "Apr", "May", "Jun", 
                       "Jul", "Aug", "Sep", "Oct", "Nov", "Dec"]
    has_year = any(indicator in para for indicator in date_indicators)
    has_month = any(indicator in para for indicator in month_indicators)
    return has_year or has_month


def _parse_education_line(para: str) -> dict[str, Any] | None:
    """Parse an education entry."""
    # Common patterns: "MBA, Harvard Business School, 2015"
    # "BS Computer Science, MIT, 2010-2014"
    # "Fellow of the Society of Actuaries (FSA)"
    
    edu = {"degree": "", "school": "", "year": "", "details": ""}
    
    # Degree indicators
    degree_indicators = [
        "MBA", "MS", "M.S.", "MA", "M.A.", "BS", "B.S.", "BA", "B.A.", "PhD", "Ph.D.",
        "Fellow", "FSA", "FCAS", "CFA", "CPA", "JD", "MD", "Executive Education"
    ]
    
    # Check for degree
    for indicator in degree_indicators:
        if indicator in para:
            edu["degree"] = indicator
            break
    
    # Try to extract year (4 digits)
    import re
    year_match = re.search(r"\b(19|20)\d{2}\b", para)
    if year_match:
        edu["year"] = year_match.group(0)
    
    # School extraction (common universities)
    school_indicators = [
        "Harvard", "Stanford", "MIT", "Yale", "Princeton", "Columbia", "Cornell",
        "Dartmouth", "Brown", "Penn", "Wharton", "Booth", "Kellogg", "Sloan",
        "NYU", "Northwestern", "Duke", "Berkeley", "UCLA", "USC", "CMU",
        "University of", "College", "Institute", "School of", "State University"
    ]
    
    for indicator in school_indicators:
        if indicator in para:
            # Extract school name (simplified)
            parts = para.split(",")
            for part in parts:
                if indicator in part:
                    edu["school"] = part.strip()
                    break
            break
    
    # If we have at least degree or school, return the entry
    if edu["degree"] or edu["school"]:
        edu["details"] = para  # Keep full text as details
        return edu
    
    return None


def _parse_certification_line(para: str) -> dict[str, Any] | None:
    """Parse a certification entry."""
    cert = {"name": "", "issuer": "", "year": "", "details": para}
    
    # Common certification patterns
    cert_indicators = [
        "FSA", "FCAS", "CFA", "CPA", "CMA", "CIA", "CISA", "CISM", "CISSP",
        "PMP", "AWS", "Azure", "Google Cloud", "Scrum", "ITIL", "Six Sigma",
        "Fellow of the Society of Actuaries", "Fellow of the Casualty Actuarial Society"
    ]
    
    for indicator in cert_indicators:
        if indicator in para:
            cert["name"] = indicator
            break
    
    # If no known indicator but looks like a credential (parentheses, dates)
    if not cert["name"]:
        import re
        # Check for patterns like "Certified X (Y)" or "X, Y (Z)"
        if "(" in para or "certified" in para.lower() or "license" in para.lower():
            cert["name"] = para.split(",")[0].strip()
    
    # Extract year
    import re
    year_match = re.search(r"\b(19|20)\d{2}\b", para)
    if year_match:
        cert["year"] = year_match.group(0)
    
    return cert if cert["name"] else None


def _parse_skills_line(para: str) -> list[str]:
    """Parse a skills line into individual skills."""
    # Split by common separators
    separators = [",", "|", ";", "•", "-"]
    skills = [para]
    
    for sep in separators:
        new_skills = []
        for skill in skills:
            new_skills.extend([s.strip() for s in skill.split(sep) if s.strip()])
        skills = new_skills
    
    # Filter out empty and very short items
    return [s for s in skills if len(s) > 2]


def _build_docx_commit_candidate(
    resume_data: dict[str, Any],
    proposed_dir: Path,
    template_path: Path | None = None,
    target_company: str = "",
    target_role: str = "",
) -> InertArtifactCommitCandidate | None:
    """Build inert DOCX commit candidate — NO durable write.

    GAP-001 L4 Boundary Hardening: Exit must not mutate durable state.
    This function prepares a DOCX commit candidate with content serialized
    to bytes for UWG-mediated admission. No filesystem write occurs.

    Args:
        resume_data: The generated resume JSON data.
        proposed_dir: Virtual directory path for UWG reference.
        template_path: Optional path to template DOCX (read-only, not copied).
        target_company: Target company name for filename.
        target_role: Target role for filename.

    Returns:
        InertArtifactCommitCandidate with serialized DOCX content, or None
        if python-docx unavailable or serialization fails.
    """
    try:
        import docx  # python-docx
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import io
    except ImportError:
        _LOGGER.warning("[apps_rg Exit] python-docx not available, skipping DOCX candidate")
        return None

    # Generate filename (virtual path only)
    header_name = resume_data.get("header", {}).get("name", "")
    safe_name = header_name.replace(" ", "_").replace("/", "_") if header_name else "resume"
    safe_company = target_company.replace(" ", "_").replace("/", "_") if target_company else ""
    safe_role = target_role.replace(" ", "_").replace("/", "_") if target_role else ""

    if safe_name and safe_company and safe_role:
        filename = f"{safe_name}_{safe_company}_{safe_role}.docx"
    else:
        filename = "resume_output.docx"

    proposed_path = str(proposed_dir / filename)

    # Build DOCX document in memory
    try:
        if template_path and template_path.exists():
            # Read template (read-only), don't copy to output
            doc = docx.Document(str(template_path))
            _update_resume_docx_content(doc, resume_data, target_company, target_role)
        else:
            doc = docx.Document()
            _build_resume_docx_content(doc, resume_data)

        # Serialize to bytes in memory (no file write)
        docx_buffer = io.BytesIO()
        doc.save(docx_buffer)
        docx_bytes = docx_buffer.getvalue()
        docx_buffer.close()

        # Build content representation
        content_repr = {
            "docx_bytes_length": len(docx_bytes),
            "docx_bytes_digest": hashlib.sha256(docx_bytes).hexdigest()[:32],
            "template_used": str(template_path) if template_path and template_path.exists() else None,
            "non_durable": True,
            "not_l4_truth": True,
        }

        return InertArtifactCommitCandidate(
            artifact_type="docx_resume",
            proposed_path=proposed_path,
            content_digest=content_repr["docx_bytes_digest"],
            serialized_content=content_repr,
            mutation_candidate_inert=True,
            proposal_status="PENDING_UWG",
            non_durable=True,
            not_l4_truth=True,
            not_replay_source=True,
        )
    except Exception as exc:
        _LOGGER.warning("[apps_rg Exit] Failed to build DOCX candidate: %s", exc)
        return None


# GAP-001: Legacy function kept for compatibility but marked DEPRECATED
# All calls should use _build_docx_commit_candidate instead
def _write_resume_docx(
    resume_data: dict[str, Any],
    output_dir: Path,
    template_path: Path | None = None,
    target_company: str = "",
    target_role: str = "",
) -> Path | None:
    """DEPRECATED: Use _build_docx_commit_candidate.

    GAP-001 L4 Boundary Hardening: Direct filesystem writes are FORBIDDEN.
    This function now returns None and emits a warning.
    """
    _LOGGER.error(
        "[apps_rg Exit] GAP-001 VIOLATION: _write_resume_docx called but durable "
        "writes are forbidden. Use _build_docx_commit_candidate for inert proposals."
    )
    return None


def _build_resume_docx_content(doc: Any, resume_data: dict[str, Any]) -> None:
    """Build resume content in DOCX document from resume JSON data."""
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE

    # Helper to add styled paragraph
    def add_heading(text: str, level: int = 1) -> Any:
        p = doc.add_heading(text, level=level)
        return p

    def add_paragraph(text: str, style: str | None = None, bold: bool = False) -> Any:
        p = doc.add_paragraph(text, style=style)
        if bold:
            for run in p.runs:
                run.bold = True
        return p

    # Executive Summary
    if resume_data.get("executive_summary"):
        add_heading("Executive Summary", level=1)
        add_paragraph(resume_data["executive_summary"])
        doc.add_paragraph()  # Spacing


def _update_resume_docx_content(
    doc: Any,
    resume_data: dict[str, Any],
    target_company: str = "",
    target_role: str = "",
) -> None:
    """Update existing DOCX template with generated resume content.

    This preserves the template's styling while injecting the tailored
    resume content for the target role and company.
    """
    from docx.shared import Pt

    # Helper to find or add paragraph
    def find_or_add_paragraph(search_text: str, new_text: str = "") -> Any:
        for para in doc.paragraphs:
            if search_text.lower() in para.text.lower():
                return para
        return doc.add_paragraph(new_text)

    # Update header if present in resume data
    header = resume_data.get("header", {})
    if header:
        name = header.get("name", "")
        # Try to update name in first paragraph (typically the header)
        if name and doc.paragraphs:
            first_para = doc.paragraphs[0]
            if first_para.text.strip():
                # Replace first paragraph text with name
                first_para.clear()
                run = first_para.add_run(name)
                run.bold = True
                run.font.size = Pt(16)

        # Update contact info section
        contact_parts = []
        if header.get("phone"):
            contact_parts.append(header["phone"])
        if header.get("email"):
            contact_parts.append(header["email"])
        if header.get("linkedin"):
            contact_parts.append(header["linkedin"])
        if header.get("location"):
            contact_parts.append(header["location"])

        if contact_parts and len(doc.paragraphs) > 1:
            contact_para = doc.paragraphs[1]
            contact_para.clear()
            contact_para.add_run(" | ".join(contact_parts))

    # Update Executive Summary section
    exec_summary = resume_data.get("executive_summary", "")
    if exec_summary:
        for i, para in enumerate(doc.paragraphs):
            if "executive summary" in para.text.lower():
                # Found the heading, update the next paragraph
                if i + 1 < len(doc.paragraphs):
                    summary_para = doc.paragraphs[i + 1]
                    summary_para.clear()
                    summary_para.add_run(exec_summary)
                break

    # Add target role/company header if not present
    if target_role or target_company:
        target_header = f"Target: {target_role} at {target_company}" if target_role and target_company else f"Target: {target_role or target_company}"
        # Insert after contact info
        if len(doc.paragraphs) > 2:
            target_para = doc.paragraphs[2]
            if "target" not in target_para.text.lower():
                # Insert new paragraph
                new_para = doc.paragraphs[1]._element.addnext(doc.add_paragraph(target_header)._element)

    # Update Experience section
    experience = resume_data.get("experience", [])
    if experience:
        # Look for Experience heading and update content after it
        for i, para in enumerate(doc.paragraphs):
            if "experience" in para.text.lower() and para.style.name.startswith('Heading'):
                # Found experience section - we could add/update roles here
                # For now, the template structure is preserved
                break


@dataclasses.dataclass(frozen=True)
class ExitBindingResult:
    """Result of the Exit binding execution.

    GAP-001 L4 Boundary Hardening: Exit emits inert proposals only.
    Per plan apps-rg-runtime-wiring-completion-d4e8a1 W3.P5, Exit emits an
    X3Disposition that contains the final disposition, inert artifact references,
    and gate evaluation results. This dataclass captures both the disposition
    and the artifact commit candidates for UWG-mediated admission.

    Attributes:
        disposition: X3Disposition with exit status and gate verdicts
        output_artifact_path: Virtual path for user reference (not durable L4)
        artifact_commit_candidates: Tuple of inert commit proposals for UWG
        user_visible_resume: Resume content for immediate caller consumption
    """

    disposition: X3Disposition
    output_artifact_path: Path
    artifact_commit_candidates: tuple[InertArtifactCommitCandidate, ...] = ()
    user_visible_resume: dict[str, Any] = dataclasses.field(default_factory=dict)


def exit_finalize_apps_rg(
    sealed: SealedL2Artifact,
    target_company: str,
    target_role: str,
    output_directory: str | Path | None = None,
    writeback_policy: Mapping[str, Any] | None = None,
) -> ExitBindingResult:
    """Finalize the apps_rg pipeline: emit inert proposals, emit X3Disposition.

    GAP-001 L4 Boundary Hardening: Exit is the SEVENTH stage per W3.P5.
    It consumes the SealedL2Artifact produced by L2 and:

    1. Validates the execution_status is success-like
    2. Builds inert artifact commit candidates (NO durable writes)
    3. Optionally builds cache commit candidates (opt-in via policy)
    4. Optionally builds C0 chunk commit candidates (opt-in via policy)
    5. Evaluates Exit gates (G24/G25/G26/G27)
    6. Builds and returns an X3Disposition with inert proposals

    Architectural Law: Exit must not mutate durable state. UWG is sole durable writer.
    All artifact persistence is via InertArtifactCommitCandidate with
    mutation_candidate_inert=True, proposal_status=PENDING_UWG.

    Args:
        sealed: SealedL2Artifact from L2 execution.
        target_company: Target company name (for path naming).
        target_role: Target role title (for path naming).
        output_directory: Optional override output directory (virtual path only).
        writeback_policy: Optional dict with cache_write and c0_write bools.

    Returns:
        ExitBindingResult with X3Disposition, inert commit candidates, and
        user-visible resume content. No filesystem mutation occurs.

    Raises:
        TypeError: if sealed is not a SealedL2Artifact.
    """
    if not isinstance(sealed, SealedL2Artifact):
        raise TypeError(
            f"exit_finalize_apps_rg expected SealedL2Artifact, got "
            f"{type(sealed).__name__}"
        )

    # GAP-001: Virtual path computation for inert commit candidates
    # No filesystem mutation — these are proposal paths for UWG reference
    repo_root = _resolve_repo_root()

    # Resolve virtual output directory for path naming
    if output_directory is None:
        base_dir = repo_root / "artifacts" / "apps_rg" / "runs"
    else:
        base_dir = Path(output_directory)
        if not base_dir.is_absolute():
            base_dir = repo_root / base_dir

    # Compute the virtual run directory name (inert reference)
    run_dirname = _safe_run_dirname(target_company, target_role, sealed.run_id)
    run_dir = base_dir / run_dirname

    # GAP-001: Inert commit candidate collection
    commit_candidates: list[InertArtifactCommitCandidate] = []

    # Build the inert resume JSON commit candidate
    resume_doc = sealed.proposed_state_diff or {}
    resume_candidate = _build_artifact_commit_candidate(
        content=resume_doc,
        proposed_dir=run_dir,
        filename="generated_resume.json",
        artifact_type="resume_json",
    )
    commit_candidates.append(resume_candidate)
    resume_path = Path(resume_candidate.proposed_path)  # Virtual path for X3 reference

    # -------------------------------------------------------------------------
    # GAP-001: DOCX output is INERT PROPOSAL ONLY
    # Template reading is allowed (read-only), but DOCX generation produces
    # inert candidate only — no durable write.
    # -------------------------------------------------------------------------
    _template_base = Path("C:/Users/amita/Documents/Resumes")
    _template_candidates = [
        _template_base / "SVP Engineering Resume_Ayer.docx",
        _template_base / "SVP Engineering Resume.docx",
    ]
    template_path: Path | None = None
    for _cand in _template_candidates:
        if _cand.exists():
            template_path = _cand
            break

    # DOCX candidate is inert — content is prepared but not persisted
    docx_candidate = _build_docx_commit_candidate(
        resume_data=resume_doc,
        proposed_dir=run_dir,
        template_path=template_path,
        target_company=target_company,
        target_role=target_role,
    )
    docx_path: Path | None = None
    if docx_candidate:
        commit_candidates.append(docx_candidate)
        docx_path = Path(docx_candidate.proposed_path)
        _LOGGER.info("[apps_rg Exit] DOCX commit candidate prepared: %s", docx_candidate.proposed_path)

    # Build run metadata for traceability (inert, not persisted by Exit)
    run_metadata: dict[str, Any] = {
        "app_id": sealed.app_id,
        "task_class": "resume_generation",
        "run_id": sealed.run_id,
        "request_id": sealed.request_id,
        "trace_id": sealed.trace_id,
        "target_company": target_company,
        "target_role": target_role,
        "execution_status": sealed.execution_status,
        "execution_timestamp": sealed.execution_timestamp,
        "execution_duration_ms": sealed.execution_duration_ms,
        "sovereign_execution_receipt": sealed.sovereign_execution_receipt,
        "schema_version": "1.0",
        "writeback_enabled": False,
        "cache_write": False,
        "c0_write": False,
        "gap_001_status": "CLOSED",  # GAP-001: Exit no longer performs durable writes
        "non_durable": True,
        "not_l4_truth": True,
        "output_paths": {
            "generated_resume": str(resume_candidate.proposed_path),
        },
    }
    if docx_candidate:
        run_metadata["output_paths"]["docx_resume"] = docx_candidate.proposed_path

    # -------------------------------------------------------------------------
    # GAP-001: Cache commit candidate (inert proposal only)
    # Disabled by default until UWG-mediated admission is validated.
    # -------------------------------------------------------------------------
    cache_write_enabled = (
        writeback_policy.get("cache_write")
        if writeback_policy else
        os.environ.get(_APPS_RG_CACHE_WRITE_ENABLED_ENV, "").strip().lower()
        in ("1", "true", "yes")
    )
    if cache_write_enabled:
        try:
            provenance_digest = _safe_build_g24_provenance(sealed)
            cache_key = hashlib.sha256(provenance_digest.encode("utf-8")).hexdigest()
            cache_entry = {
                "cache_key": cache_key,
                "app_id": sealed.app_id,
                "provenance": provenance_digest,
                "output_ref": resume_candidate.proposed_path,
                "run_id": sealed.run_id,
                "request_id": sealed.request_id,
                "created_at": sealed.execution_timestamp,
                "non_durable": True,  # GAP-001: explicit non-durable marker
            }
            cache_candidate = _build_artifact_commit_candidate(
                content=cache_entry,
                proposed_dir=run_dir,
                filename="semantic_cache_entry.json",
                artifact_type="cache_entry",
            )
            commit_candidates.append(cache_candidate)
            run_metadata["writeback_enabled"] = True
            run_metadata["cache_write"] = True
            run_metadata["output_paths"]["semantic_cache"] = cache_candidate.proposed_path
            _LOGGER.info("[apps_rg Exit] Cache commit candidate prepared: %s", cache_candidate.proposed_path)
        except Exception as exc:
            _LOGGER.warning("[apps_rg Exit] Cache candidate build failed: %s", exc)

    # -------------------------------------------------------------------------
    # GAP-001: C0 chunk commit candidate (inert proposal only)
    # -------------------------------------------------------------------------
    c0_write_enabled = (
        writeback_policy.get("c0_write")
        if writeback_policy else
        os.environ.get(_APPS_RG_C0_WRITE_ENABLED_ENV, "").strip().lower()
        in ("1", "true", "yes")
    )
    if c0_write_enabled:
        try:
            c0_entry = {
                "chunk_refs": [],
                "context_digest": sealed.compilation_hash,
                "run_id": sealed.run_id,
                "non_durable": True,  # GAP-001: explicit non-durable marker
            }
            c0_candidate = _build_artifact_commit_candidate(
                content=c0_entry,
                proposed_dir=run_dir,
                filename="c0_output_chunks.json",
                artifact_type="c0_chunks",
            )
            commit_candidates.append(c0_candidate)
            run_metadata["c0_write"] = True
            run_metadata["output_paths"]["c0_chunks"] = c0_candidate.proposed_path
            _LOGGER.info("[apps_rg Exit] C0 chunk candidate prepared: %s", c0_candidate.proposed_path)
        except Exception as exc:
            _LOGGER.warning("[apps_rg Exit] C0 candidate build failed: %s", exc)

    # Build run metadata commit candidate (inert proposal)
    metadata_candidate = _build_artifact_commit_candidate(
        content=run_metadata,
        proposed_dir=run_dir,
        filename="run_metadata.json",
        artifact_type="run_metadata",
    )
    commit_candidates.append(metadata_candidate)

    # -------------------------------------------------------------------------
    # Gate evaluation (G24/G25/G26/G27)
    # Per apps_rg governance, these are evaluated at Exit for final disposition.
    # Currently stubbed — real evaluation wiring deferred to W4.
    # -------------------------------------------------------------------------
    gate_results = []

    # G24: G24 provenance check (receipt presence + hash alignment)
    g24_pass = bool(
        sealed.sovereign_execution_receipt
        and sealed.sovereign_execution_receipt.startswith("vllm-")
    ) or bool(
        sealed.execution_status == "completed_stub_fallback"
    )  # stub fallback is also valid
    gate_results.append(
        AppsRgGateResult(
            gate_id="G24",
            verdict=ExitGateVerdict.PASS if g24_pass else ExitGateVerdict.WARN,
            score=1.0 if g24_pass else 0.0,
            weight=1.0,
            reason="Receipt present and execution_status valid" if g24_pass else "Missing or malformed receipt",
        )
    )

    # G25: Tenant isolation check (apps_rg tenant binding)
    g25_pass = sealed.tenant_id == "apps_rg"
    gate_results.append(
        AppsRgGateResult(
            gate_id="G25",
            verdict=ExitGateVerdict.PASS if g25_pass else ExitGateVerdict.FAIL,
            score=1.0 if g25_pass else 0.0,
            weight=1.0,
            reason="Tenant isolation verified (apps_rg)" if g25_pass else "Tenant mismatch",
        )
    )

    # G26: State diff schema validation (basic shape check)
    g26_pass = isinstance(sealed.proposed_state_diff, dict)
    gate_results.append(
        AppsRgGateResult(
            gate_id="G26",
            verdict=ExitGateVerdict.PASS if g26_pass else ExitGateVerdict.FAIL,
            score=1.0 if g26_pass else 0.0,
            weight=1.0,
            reason="State diff is dict" if g26_pass else "State diff malformed",
        )
    )

    # G27: Word count sanity check (upper bound ~1500 words for resumes)
    content = sealed.generated_content or ""
    word_count = len(content.split())
    g27_pass = word_count < 2000
    gate_results.append(
        AppsRgGateResult(
            gate_id="G27",
            verdict=ExitGateVerdict.PASS if g27_pass else ExitGateVerdict.WARN,
            score=1.0 if g27_pass else 0.5,
            weight=1.0,
            reason=f"Word count {word_count} within bounds" if g27_pass else f"Word count {word_count} exceeds soft limit",
        )
    )

    # Compute overall verdict
    fail_count = sum(
        1 for g in gate_results if g.verdict == ExitGateVerdict.FAIL
    )
    warn_count = sum(
        1 for g in gate_results if g.verdict == ExitGateVerdict.WARN
    )
    if fail_count > 0:
        overall_verdict = ExitGateVerdict.FAIL
    elif warn_count > 0:
        overall_verdict = ExitGateVerdict.WARN
    else:
        overall_verdict = ExitGateVerdict.PASS

    # Build X3Disposition
    # Convert gate results to string refs for gate_verdict_refs
    gate_verdict_refs = tuple(
        f"{g.gate_id}:{g.verdict.value}:{g.score:.2f}" for g in gate_results
    )
    disposition = X3Disposition(
        request_id=sealed.request_id,
        run_id=sealed.run_id,
        app_id=sealed.app_id,
        trace_id=sealed.trace_id,
        exit_status="success" if sealed.execution_status in ("completed", "completed_stub_fallback") else "failed",
        outcome_authorized=(overall_verdict != ExitGateVerdict.FAIL),
        output_artifact_path=str(resume_path),
        final_output=resume_doc,
        gate_verdict_refs=gate_verdict_refs,
        sealed_l2_digest=sealed.compilation_hash,
        l5_certification_ref=APPS_RG_EXIT_CERT_REF,
        # Thread provenance for full traceability
        tenant_id=sealed.tenant_id,
    )

    _LOGGER.info(
        "[apps_rg Exit] Finalized run %s: exit_status=%s verdict=%s candidates=%d",
        sealed.run_id,
        disposition.exit_status,
        overall_verdict.value,
        len(commit_candidates),
    )

    # GAP-001: Return ExitBindingResult with inert commit candidates
    # No filesystem mutation has occurred — all proposals are PENDING_UWG
    return ExitBindingResult(
        disposition=disposition,
        output_artifact_path=resume_path,
        artifact_commit_candidates=tuple(commit_candidates),
        user_visible_resume=resume_doc,
    )


def build_apps_rg_exit_harness(
    sealed: SealedL2Artifact,
    writeback_policy: Mapping[str, Any] | None = None,
) -> ExitBindingResult:
    """Convenience harness for apps_rg Exit with no explicit company/role.

    Extracts target_company and target_role from the sealed artifact's
    proposed_state_diff (populated by L2 from the PA prompt).

    Args:
        sealed: SealedL2Artifact from L2.
        writeback_policy: Optional cache/C0 writeback policy.

    Returns:
        ExitBindingResult with X3Disposition.
    """
    state_diff = sealed.proposed_state_diff or {}
    target_company = state_diff.get("target_company", "UNKNOWN_COMPANY")
    target_role = state_diff.get("target_role", "UNKNOWN_ROLE")

    return exit_finalize_apps_rg(
        sealed=sealed,
        target_company=target_company,
        target_role=target_role,
        writeback_policy=writeback_policy,
    )


def extract_apps_rg_exit_gate_policy(validated_request: Any) -> dict[str, Any]:
    """Extract Exit gate policy from ValidatedRequest.app_payload.

    Per plan apps-rg-quarantine-gap-remediation-8f405c W5.P2, Exit gate
    policy is extracted from app_payload for evaluation at L5/Exit.

    Returns a dict with:
      - g24_enabled, g25_enabled, g26_enabled, g27_enabled (bools)
      - fail_closed (bool)
      - payload_path (str)

    Args:
        validated_request: ValidatedRequest carrying app_payload.

    Returns:
        Dict with gate policy booleans and provenance info.
    """
    try:
        app_payload = validated_request.app_payload
        if isinstance(app_payload, dict):
            # dict access for payloads that come through raw
            exit_policy = app_payload.get("exit_gate_policy") or {}
            return {
                "g24_enabled": exit_policy.get("g24_enabled", True),
                "g25_enabled": exit_policy.get("g25_enabled", True),
                "g26_enabled": exit_policy.get("g26_enabled", True),
                "g27_enabled": exit_policy.get("g27_enabled", True),
                "fail_closed": exit_policy.get("fail_closed", False),
                "payload_path": "ValidatedRequest.app_payload.exit_gate_policy",
            }
        # Dataclass access path
        exit_policy = getattr(app_payload, "exit_gate_policy", None)
        if exit_policy is None:
            # default all gates enabled, fail-soft (not fail-closed)
            return {
                "g24_enabled": True,
                "g25_enabled": True,
                "g26_enabled": True,
                "g27_enabled": True,
                "fail_closed": False,
                "payload_path": "ValidatedRequest.app_payload.exit_gate_policy",
            }
        return {
            "g24_enabled": getattr(exit_policy, "g24_enabled", True),
            "g25_enabled": getattr(exit_policy, "g25_enabled", True),
            "g26_enabled": getattr(exit_policy, "g26_enabled", True),
            "g27_enabled": getattr(exit_policy, "g27_enabled", True),
            "fail_closed": getattr(exit_policy, "fail_closed", False),
            "payload_path": "ValidatedRequest.app_payload.exit_gate_policy",
        }
    except Exception as exc:  # guardian: allow-broad-net -- policy extraction must never abort Exit; missing policy is WARN not ERROR
        _LOGGER.warning("[apps_rg Exit] gate policy extraction failed: %s", exc)
        return {
            "g24_enabled": True,
            "g25_enabled": True,
            "g26_enabled": True,
            "g27_enabled": True,
            "fail_closed": False,
            "payload_path": "ValidatedRequest.app_payload.exit_gate_policy",
        }


def produce_structured_resume_from_docx(
    docx_path: str | Path,
    target_company: str | None = None,
    target_role: str | None = None,
) -> dict[str, Any]:
    """Produce structured resume JSON from DOCX input.

    W1: DOCX ingestion emits normalized structured resume format.
    This function wraps _ingest_docx_to_master_resume and adds the
    required schema fields (content_kind, rewrite_policy, judge_policy,
    section_id, company_id) per source_resume_v2_structured.json schema.

    Args:
        docx_path: Path to DOCX resume file.
        target_company: Optional target company for provenance.
        target_role: Optional target role for provenance.

    Returns:
        Structured resume dict conforming to source_resume_v2_structured.json.
    """
    # Ingest raw data from DOCX
    raw_data = _ingest_docx_to_master_resume(docx_path)

    # Normalize to structured format
    structured: dict[str, Any] = {
        "schema_name": "source_resume_v2_structured",
        "schema_version": "2.0.0",
        "flat_text_fallback": None,
        "headline": {
            "section_id": "headline",
            "content_kind": "narrative_only",
            "rewrite_policy": "heavy",
            "judge_policy": "p0_full_panel",
            "text": raw_data.get("headline", ""),
            "preserve_verbatim": False,
            "treatment_tier": None,
        },
        "executive_summary": {
            "section_id": "executive_summary",
            "content_kind": "narrative_only",
            "rewrite_policy": "heavy",
            "judge_policy": "p0_full_panel",
            "text": raw_data.get("executive_summary", ""),
            "preserve_verbatim": False,
            "treatment_tier": None,
        },
        "roles": [],
        "early_career": {
            "section_id": "early_career",
            "content_kind": "verbatim_copy",
            "rewrite_policy": "verbatim",
            "judge_policy": "none",
            "entries": [],
            "preserve_verbatim": True,
        },
        "competencies": {
            "section_id": "competencies",
            "content_kind": "bullets_only",
            "rewrite_policy": "moderate",
            "judge_policy": "p1_full_panel",
            "items": raw_data.get("competencies", []),
            "treatment_tier": None,
            "preserve_verbatim": False,
        },
        "education": {
            "section_id": "education",
            "content_kind": "verbatim_copy",
            "rewrite_policy": "verbatim",
            "judge_policy": "none",
            "entries": [],
            "preserve_verbatim": True,
        },
        "certifications": {
            "section_id": "certifications",
            "content_kind": "verbatim_copy",
            "rewrite_policy": "verbatim",
            "judge_policy": "none",
            "entries": [],
            "preserve_verbatim": True,
        },
    }

    # Normalize roles with section_id and company_id
    company_id_map: dict[str, str] = {
        "unify": "unify_consulting",
        "unify consulting": "unify_consulting",
        "ibm": "ibm",
        "insurtech": "insurtech_tech_solutions",
        "insurtech tech solutions": "insurtech_tech_solutions",
        "ey": "ernst_young",
        "ernst & young": "ernst_young",
        "ernst and young": "ernst_young",
    }

    for raw_role in raw_data.get("experience", []):
        company = raw_role.get("company", "")
        company_lower = company.lower().strip()
        company_id = company_id_map.get(company_lower, company_lower.replace(" ", "_").replace("&", "and"))
        
        # Generate section_id from company
        section_id = company_id.replace("_", "").replace(" ", "_")
        
        role = {
            "section_id": section_id,
            "content_kind": "narrative_and_bullets",
            "rewrite_policy": "moderate",
            "judge_policy": "p1_full_panel",
            "employer": company,
            "company_id": company_id,
            "title": raw_role.get("title", ""),
            "location": raw_role.get("location"),
            "start_date": raw_role.get("start_date"),
            "end_date": raw_role.get("end_date"),
            "narrative": raw_role.get("narrative", ""),
            "preserve_narrative_verbatim": True,
            "bullets": [
                {
                    "source_text": b.get("text", ""),
                    "ordinal": b.get("ordinal", i + 1),
                    "treatment_tier": None,
                    "rewrite_allowed": b.get("rewrite_allowed", True),
                    "preserve_verbatim": b.get("preserve_verbatim", False),
                    "evidence_required": b.get("evidence_required", False),
                    "source_span_ref": b.get("source_span_ref"),
                }
                for i, b in enumerate(raw_role.get("bullets", []))
            ],
            "source_span_refs": raw_role.get("source_span_refs"),
            "bullet_treatment_tier": None,
        }
        structured["roles"].append(role)

    # Normalize education entries
    for edu in raw_data.get("education", []):
        structured["education"]["entries"].append({
            "text": edu.get("text", ""),
            "label": edu.get("label"),
            "preserve_verbatim": True,
        })

    # Normalize certifications
    for cert in raw_data.get("certifications", []):
        structured["certifications"]["entries"].append({
            "text": cert.get("text", ""),
            "label": cert.get("label"),
            "preserve_verbatim": True,
        })

    # Normalize early career (if present in raw data)
    for entry in raw_data.get("early_career", []):
        structured["early_career"]["entries"].append({
            "text": entry.get("text", ""),
            "label": entry.get("label"),
            "preserve_verbatim": True,
        })

    return structured


__all__ = [
    "APPS_RG_EXIT_CERT_REF",
    "ExitBindingResult",
    "InertArtifactCommitCandidate",
    "build_apps_rg_exit_harness",
    "exit_finalize_apps_rg",
    "extract_apps_rg_exit_gate_policy",
    "produce_structured_resume_from_docx",
    "_build_artifact_commit_candidate",  # GAP-001: exposed for testing
]

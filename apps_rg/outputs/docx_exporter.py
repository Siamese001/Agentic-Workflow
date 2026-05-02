"""DOCX exporter for apps_rg generated_resume.json.

Converts the structured resume artifact to a recruiter-ready Word document
by CLONING the user's base SVP resume template and substituting content.
This guarantees the output matches the user's exact font, spacing, section
headings, and "bold-first" style so the generated DOCX can be used directly
without manual reformatting.

Template: apps_shared/templates/svp_resume_template.docx
  (Copy of "SVP Engineering Resume_Ayer.docx" — recruiter-approved layout)

Layout matched to template:
  1. Name (Heading 1, centered)
  2. Tagline (12pt bold, centered)         <- HOP-4A headline
  3. Contact (10pt, centered)
  4. EXECUTIVE SUMMARY  (Heading 2)        <- HOP-4B
  5. PROFESSIONAL EXPERIENCE (Heading 2)
       per role:
         "Company | Location"              <- Company bold, 12pt
         "Title | Dates"                   <- Title bold
         context (Normal)
         "Label: Body" bullets             <- Label bold
  6. EDUCATION (Heading 2)
  7. CERTIFICATIONS & CREDENTIALS (Heading 2)
  8. ENGINEERING & PLATFORM COMPETENCIES (Heading 2)
       "Category: Terms" lines             <- Category bold

Design notes:
- python-docx (1.2+) only. No pandoc, no headless Chrome, no LaTeX.
- Deterministic: identical input -> identical output bytes.
- Honors the duplicate-section schema by preferring ``professional_experience``
  (richer: has bullet_pool, highlights) and falling back to ``experience``.
"""
from __future__ import annotations

import argparse
import json
import re
import sys
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Inches


# ---------------------------------------------------------------------------
# Style helpers
# ---------------------------------------------------------------------------

ACCENT = RGBColor(0x1F, 0x3A, 0x5F)  # deep navy — executive tone
GREY = RGBColor(0x55, 0x55, 0x55)


def _strip_md_bold(text: str) -> tuple[str, str]:
    """Split a string of the form '**bold:** rest' into (bold, rest).

    Returns ('', text) if no bold prefix is present.
    """
    m = re.match(r"^\s*\*\*([^*]+)\*\*[:\s]*(.*)$", text)
    if m:
        # Strip trailing punctuation from label so we don't render "Label::"
        label = m.group(1).strip().rstrip(":;.,—-")
        return label, m.group(2).strip()
    return "", text.strip()


def _strip_bullet_glyph(text: str) -> str:
    """Remove leading bullet glyphs (•, ·, –, —, dash) and surrounding whitespace.

    Does NOT strip leading ``*`` because that conflicts with markdown bold
    markers handled by :func:`_strip_md_bold`.
    """
    return re.sub(r"^[\s\u2022\u00b7\u2013\u2014\-]+", "", text).strip()


def _add_horizontal_rule(paragraph) -> None:
    """Insert a thin horizontal rule under a paragraph (used after section heads)."""
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1F3A5F")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def _section_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text.upper())
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = ACCENT
    _add_horizontal_rule(p)


def _bullet(doc: Document, text: str, bold_prefix: str = "") -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(0.25)
    if bold_prefix:
        b = p.add_run(f"{bold_prefix}: ")
        b.font.bold = True
        b.font.size = Pt(10.5)
    rest = p.add_run(text)
    rest.font.size = Pt(10.5)


def _kv_line(doc: Document, label: str, value: str, *, bold_label: bool = True) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    if bold_label:
        run = p.add_run(f"{label}: ")
        run.font.bold = True
        run.font.size = Pt(10)
    p.add_run(value).font.size = Pt(10)


# ---------------------------------------------------------------------------
# Data extraction
# ---------------------------------------------------------------------------


def _bullet_text(b: Any) -> str:
    if isinstance(b, str):
        return b
    if isinstance(b, dict):
        return b.get("bullet_text") or b.get("text") or ""
    return str(b)


def _format_role_label(company: str, title: str | None, dates: dict, location: str | None) -> tuple[str, str]:
    """Produce (line1, line2) for a role header.

    line1 = "Title — Company" (bold)
    line2 = "Location  |  Mon YYYY – Present" (italic grey)
    """
    title = (title or "").strip()
    company = (company or "").strip()
    if title and title.lower() not in company.lower():
        line1 = f"{title} — {company}"
    else:
        line1 = company or title

    start = (dates or {}).get("start") or ""
    end = (dates or {}).get("end") or "Present"
    span = " – ".join(x for x in (start, end) if x)
    parts = [p for p in ((location or "").strip(), span) if p]
    line2 = "  |  ".join(parts)
    return line1, line2


def _experience(resume: dict) -> list[dict]:
    """Pick the richer of `professional_experience` or `experience`."""
    pe = resume.get("professional_experience") or []
    ex = resume.get("experience") or []
    return pe if pe else ex


_YEAR_RE = re.compile(r"\b(19|20)\d{2}\b")


def _is_early_career_role(role: dict) -> bool:
    """Match narrative_pass._is_early_career — end_year < 2010.

    Handles `dates: {end: 'YYYY'|'Present'|'Month YYYY'}` and falls back
    to title markers like 'early career' when no date info is present.
    """
    import datetime as _dt

    dates = role.get("dates") or {}
    end_year = None
    direct = role.get("end_year") or role.get("end")
    try:
        if direct is not None:
            end_year = int(direct)
    except (TypeError, ValueError):
        end_year = None
    if end_year is None and isinstance(dates, dict):
        end = dates.get("end")
        if isinstance(end, str):
            if end.strip().lower() in {"present", "current", "now"}:
                end_year = _dt.date.today().year
            else:
                m = _YEAR_RE.search(end)
                if m:
                    end_year = int(m.group(0))
    if end_year is not None:
        return end_year < 2010
    blob = " ".join(
        str(role.get(k) or "") for k in ("name", "company", "title")
    ).lower()
    return "early career" in blob or "earlier career" in blob


# ---------------------------------------------------------------------------
# Builder
# ---------------------------------------------------------------------------


_REPO_ROOT = Path(__file__).resolve().parents[2]
_DEFAULT_TEMPLATE = _REPO_ROOT / "apps_shared" / "templates" / "svp_resume_template.docx"


def _clone_template(template_path: Path) -> Document:
    """Load the template and strip its body content while preserving styles.

    The template's `styles.xml` (fonts, sizes, indents, colors) stays intact
    — we only remove the paragraphs/tables. This guarantees the generated
    document inherits the exact look-and-feel of the recruiter-approved
    base resume.
    """
    doc = Document(str(template_path))
    body = doc.element.body
    # Remove every paragraph and table child. Keep sectPr (page setup).
    for child in list(body):
        if child.tag == qn("w:sectPr"):
            continue
        body.remove(child)
    return doc


def _add_para(
    doc: Document,
    *,
    style: str | None = None,
    align=None,
    space_before: float | None = None,
    space_after: float | None = None,
    left_indent: float | None = None,
):
    p = doc.add_paragraph()
    if style:
        try:
            p.style = doc.styles[style]
        except KeyError:
            pass  # fall back to default
    if align is not None:
        p.alignment = align
    fmt = p.paragraph_format
    if space_before is not None:
        fmt.space_before = Pt(space_before)
    if space_after is not None:
        fmt.space_after = Pt(space_after)
    if left_indent is not None:
        fmt.left_indent = Pt(left_indent)
    return p


def _add_run(para, text: str, *, bold: bool | None = None, size: float | None = None, italic: bool | None = None):
    r = para.add_run(text)
    if bold is not None:
        r.bold = bold
    if italic is not None:
        r.italic = italic
    if size is not None:
        r.font.size = Pt(size)
    return r


def _bold_prefixed(para, text: str) -> None:
    """Render 'Label: Body' with the label bold and the body normal.

    Falls back to a single non-bold run when no label is detected.
    """
    # Extract explicit **Label:** markdown if present.
    label, body = _strip_md_bold(text)
    if not label:
        # Infer "Label: body" split when text contains a short prefix ending in ':'
        m = re.match(r"^\s*([A-Z][A-Za-z0-9&,\- ]{1,60}):\s+(.*)$", text.strip())
        if m:
            label = m.group(1).strip()
            body = m.group(2).strip()
        else:
            _add_run(para, text.strip())
            return
    _add_run(para, f"{label}:", bold=True)
    if body:
        _add_run(para, f" {body}")


def _format_dates(dates: dict) -> str:
    if not isinstance(dates, dict):
        return ""
    start = dates.get("start", "")
    end = dates.get("end", "")
    if start and end:
        return f"{start} – {end}"
    return str(start or end or "")


def build_docx(
    resume: dict,
    *,
    target_role: str | None = None,
    target_company: str | None = None,
    template_path: Path | None = None,
) -> Document:
    """Build the DOCX by cloning the user's recruiter-approved template.

    Section order matches the template exactly:
      Name -> Tagline (HOP-4A headline) -> Contact -> SUMMARY -> EXPERIENCE
      -> EDUCATION -> CERTS -> COMPETENCIES.

    Bold-first discipline matches the template:
      - Company bold on "Company | Location" line
      - Title bold on "Title | Dates" line
      - Label bold on "Label: Body" bullets and competencies
    """
    _ = (target_role, target_company)  # kept for CLI signature compat
    tmpl = Path(template_path) if template_path else _DEFAULT_TEMPLATE
    if tmpl.exists():
        doc = _clone_template(tmpl)
    else:
        # Graceful fallback to an empty doc (tests / smoke environments).
        doc = Document()
        for section in doc.sections:
            section.top_margin = Inches(0.6)
            section.bottom_margin = Inches(0.6)
            section.left_margin = Inches(0.7)
            section.right_margin = Inches(0.7)

    owner = resume.get("owner") or {}
    name = owner.get("name") or "Candidate"
    contact = owner.get("contact") or {}

    # ---- Name (Heading 1) ----
    h = _add_para(doc, style="Heading 1", align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_run(h, name)

    # ---- Tagline — HOP-4A headline, centered, 12pt bold ----
    headline = (resume.get("headline") or "").strip()
    if headline:
        t = _add_para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
        _add_run(t, headline, bold=True, size=12)

    # ---- Contact line ----
    parts: list[str] = []
    if contact.get("phone"):
        parts.append(contact["phone"])
    if contact.get("email"):
        parts.append(contact["email"])
    if contact.get("linkedin"):
        parts.append(contact["linkedin"].replace("https://www.", "").replace("https://", ""))
    if contact.get("github"):
        parts.append(contact["github"].replace("https://www.", "").replace("https://", ""))
    if contact.get("location"):
        parts.append(contact["location"])
    if parts:
        c = _add_para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
        # Use bold pipes for visual separation (matches template).
        for i, piece in enumerate(parts):
            if i > 0:
                _add_run(c, " ", size=10)
                _add_run(c, "|", bold=True, size=10)
                _add_run(c, " ", size=10)
            _add_run(c, str(piece), size=10)

    # ---- EXECUTIVE SUMMARY (Heading 2) ----
    summary = (
        resume.get("executive_summary")
        or resume.get("professional_summary")
        or resume.get("summary")
    )
    if summary:
        _add_run(_add_para(doc, style="Heading 2"), "EXECUTIVE SUMMARY")
        p = _add_para(doc, space_after=4)
        _add_run(p, summary if isinstance(summary, str) else json.dumps(summary))

    # ---- PROFESSIONAL EXPERIENCE (Heading 2) ----
    exp = _experience(resume)
    if exp:
        _add_run(_add_para(doc, style="Heading 2"), "PROFESSIONAL EXPERIENCE")
        for role in exp:
            if _is_early_career_role(role):
                continue

            company = role.get("company", "").strip()
            location = (role.get("location") or "").strip()
            title = (role.get("title") or "").strip()
            date_str = _format_dates(role.get("dates") or {})

            # Line 1: "Company | Location" — Company bold, 12pt
            p1 = _add_para(doc, space_before=6, space_after=0)
            _add_run(p1, company, bold=True, size=12)
            if location:
                _add_run(p1, " ", size=12)
                _add_run(p1, "|", bold=True, size=12)
                _add_run(p1, " ", size=12)
                _add_run(p1, location, size=12)

            # Line 2: "Title | Dates" — Title bold
            if title or date_str:
                p2 = _add_para(doc, space_after=0)
                if title:
                    _add_run(p2, title, bold=True)
                if title and date_str:
                    _add_run(p2, " ")
                    _add_run(p2, "|", bold=True)
                    _add_run(p2, " ")
                if date_str:
                    _add_run(p2, date_str)

            # Context line (Normal, non-bold)
            context = role.get("context")
            if isinstance(context, str) and context.strip():
                pc = _add_para(doc, space_before=2, space_after=2)
                _add_run(pc, context.strip())

            # Bullets — "Label: Body" with label bold
            role_bullets = (
                role.get("bullets")
                or role.get("bullet_pool")
                or role.get("highlights")
                or []
            )
            for b in role_bullets:
                if isinstance(b, dict):
                    label = (b.get("label") or "").strip()
                    body = (b.get("text") or b.get("primary") or "").strip()
                    text = f"{label}: {body}" if label else body
                else:
                    text = _bullet_text(b)
                if not text:
                    continue
                pb = _add_para(doc, style="List Paragraph", space_before=3)
                _bold_prefixed(pb, text)

        # Early career compressed line (HOP-4H) — italic Normal
        early = resume.get("early_career")
        if early and isinstance(early, str):
            pe = _add_para(doc, space_before=6, space_after=2)
            _add_run(pe, early.strip(), italic=True)

    # ---- EDUCATION (Heading 2) ----
    edu = resume.get("education") or []
    if edu:
        _add_run(_add_para(doc, style="Heading 2"), "EDUCATION")
        for e in edu:
            p = _add_para(doc, style="Heading 3")
            if isinstance(e, str):
                _add_run(p, e)
                continue
            degree = (e.get("degree") or "").strip()
            inst = (e.get("institution") or "").strip()
            notes = (e.get("notes") or "").strip()
            line = f"{degree}, {inst}" if degree and inst else (degree or inst)
            if notes:
                line = f"{line} ({notes})"
            _add_run(p, line)

    # ---- CERTIFICATIONS & CREDENTIALS (Heading 2) ----
    certs = resume.get("certifications_and_credentials") or []
    if certs:
        _add_run(_add_para(doc, style="Heading 2"), "CERTIFICATIONS & CREDENTIALS")
        for c in certs:
            if isinstance(c, str):
                text = c
            elif isinstance(c, dict):
                n = str(c.get("name") or c.get("text") or "").strip()
                issuer = str(c.get("issuer") or "").strip()
                year = c.get("year") or c.get("date") or ""
                if n and issuer and year:
                    text = f"{n}, {issuer}, {year}"
                elif n and issuer:
                    text = f"{n}, {issuer}"
                elif n:
                    text = n
                else:
                    text = ""
            else:
                text = ""
            if text:
                p = _add_para(doc, style="Heading 3")
                _add_run(p, text)

    # ---- ENGINEERING & PLATFORM COMPETENCIES (Heading 2, at bottom) ----
    # Match the template's section title + position. Reads the SVP master's
    # engineering_and_platform_competencies when present; falls back to the
    # HOP-4C `competencies` list.
    eng_comps = resume.get("engineering_and_platform_competencies") or []
    comps = resume.get("competencies") or resume.get("core_competencies") or []
    if eng_comps:
        _add_run(_add_para(doc, style="Heading 2"), "ENGINEERING & PLATFORM COMPETENCIES")
        for entry in eng_comps:
            if isinstance(entry, dict):
                area = (entry.get("area") or entry.get("name") or "").strip()
                skills = (entry.get("skills") or entry.get("items") or "").strip()
                if not (area or skills):
                    continue
                line = f"{area}: {skills}" if area and skills else (area or skills)
            else:
                line = str(entry).strip()
            if not line:
                continue
            p = _add_para(doc, space_before=6, left_indent=3.6)
            _bold_prefixed(p, line)
    elif comps:
        _add_run(_add_para(doc, style="Heading 2"), "ENGINEERING & PLATFORM COMPETENCIES")
        for entry in comps:
            text = entry if isinstance(entry, str) else entry.get("text", str(entry))
            text = _strip_bullet_glyph(text).strip()
            if not text:
                continue
            p = _add_para(doc, space_before=6, left_indent=3.6)
            _bold_prefixed(p, text)

    return doc


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------


def _latest_run_dir(repo_root: Path) -> Path:
    base = repo_root / "artifacts" / "apps_rg" / "runs"
    if not base.exists():
        sys.exit(f"No apps_rg runs at {base}")
    runs = sorted(
        (p for p in base.iterdir() if p.is_dir() and p.name[:8].isdigit()),
        key=lambda p: p.name,
        reverse=True,
    )
    if not runs:
        sys.exit(f"No timestamped run dirs under {base}")
    return runs[0]


def _slugify(s: str) -> str:
    return re.sub(r"[^A-Za-z0-9]+", "_", s).strip("_")


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description=__doc__.splitlines()[0])
    parser.add_argument("--run-dir", type=Path, default=None, help="apps_rg run directory")
    parser.add_argument("--input", type=Path, default=None, help="Path to generated_resume.json (overrides --run-dir)")
    parser.add_argument("--output", type=Path, default=None, help="Output .docx path")
    parser.add_argument("--target-name", type=str, default=None, help="Output filename stem (e.g. 'Amit_Ayer_Resume_Blend360_SVP')")
    parser.add_argument("--target-role", type=str, default=None, help="Optional target role for header tag")
    parser.add_argument("--target-company", type=str, default=None, help="Optional target company for header tag")
    args = parser.parse_args(argv)

    repo_root = Path(__file__).resolve().parents[2]

    if args.input:
        run_dir = args.input.parent
        json_path = args.input
    else:
        run_dir = args.run_dir or _latest_run_dir(repo_root)
        json_path = run_dir / "generated_resume.json"

    if not json_path.exists():
        sys.exit(f"Missing {json_path}")

    resume = json.loads(json_path.read_text(encoding="utf-8"))
    name = (resume.get("owner") or {}).get("name", "Candidate")

    doc = build_docx(
        resume,
        target_role=args.target_role,
        target_company=args.target_company,
    )

    if args.output:
        out = args.output
    else:
        stem = args.target_name or f"{_slugify(name)}_Resume"
        out = run_dir / f"{stem}.docx"

    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    print(f"Wrote {out}  ({out.stat().st_size:,} bytes)")
    return 0


if __name__ == "__main__":
    sys.exit(main())

"""Export an assembled final_resume.json (spine aggregation output) to a polished .docx.

The integrated runtime is JSON-only by design (plan apps-rg-docx-output-removal-4650ff);
this operator tool renders the canonical product artifact on demand. It walks
final_resume.json with the SAME accessors the full-resume judge flattener uses
(apps_rg.runtime.assembly.full_resume_text), so the document matches the judged text
exactly — no content is invented or reordered at export time.

Usage:
    python ops_scripts/apps_rg/export_final_resume_docx.py <run_dir_or_final_resume.json> [out.docx]

The legacy ops_scripts/apps_rg/export_to_docx.py consumes the retired pre-spine
generated_resume_<ts>.json shape and is unrelated.
"""
from __future__ import annotations

import json
import sys
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

REPO_ROOT = Path(__file__).resolve().parents[2]
sys.path.insert(0, str(REPO_ROOT))

from apps_rg.runtime.assembly.full_resume_text import (  # noqa: E402
    CERTIFICATIONS_AND_CREDENTIALS_HEADING,
    ENGINEERING_PLATFORM_COMPETENCIES_HEADING,
    base_role_headers_from_final_resume,
    bullets_from_list,
    contact_line,
    exp_header,
    format_cert,
    format_edu,
)

ACCENT = RGBColor(0x1F, 0x3A, 0x5F)


def _is_spine_shaped(path: Path) -> bool:
    """True when the JSON is the spine-aggregation product this builder renders.

    The spine ``final_resume.json`` (under ``final_resume_assembly/``) carries
    ``candidate_identity`` + a *list* of ``sections`` with ``l2_output_snapshot``.
    The flat ``outputs/final_resume.json`` export sibling has a different schema
    (top-level ``candidate_name``, ``sections`` as a *dict*) that this builder
    cannot render — picking it silently yields a hollow DOCX.
    """
    try:
        d = json.loads(path.read_text(encoding="utf-8"))
    except (OSError, ValueError):
        return False
    return isinstance(d.get("candidate_identity"), dict) and isinstance(d.get("sections"), list)


def _resolve_final_resume(arg: str) -> Path:
    p = Path(arg)
    if p.is_file():
        return p
    candidates = sorted(p.rglob("final_resume.json"))
    if not candidates:
        raise FileNotFoundError(f"No final_resume.json under {p}")
    # Prefer the spine-aggregation product over the flat ``outputs/`` export
    # sibling (which sorts last alphabetically and would otherwise win via
    # ``[-1]``, then render empty). Among spine-shaped files prefer the one
    # under ``final_resume_assembly/``.
    spine = [c for c in candidates if _is_spine_shaped(c)]
    if spine:
        for c in spine:
            if c.parent.name == "final_resume_assembly":
                return c
        return spine[0]
    raise ValueError(
        "No spine-shaped final_resume.json (candidate_identity + list sections) "
        f"under {p}; found {[str(c) for c in candidates]}. "
        "Point at final_resume_assembly/final_resume.json."
    )


def _heading(doc: Document, text: str) -> None:
    para = doc.add_paragraph()
    run = para.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = ACCENT
    para.paragraph_format.space_before = Pt(10)
    para.paragraph_format.space_after = Pt(3)


def _para(doc: Document, text: str, *, bold: bool = False, italic: bool = False,
          size: int = 10, center: bool = False, space_after: int = 2) -> None:
    para = doc.add_paragraph()
    if center:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    para.paragraph_format.space_after = Pt(space_after)


def _bullets(doc: Document, snap: dict[str, Any]) -> None:
    for line in bullets_from_list(snap.get("bullets") or []):
        _para(doc, line, size=10)


def _gap(doc: Document, section_id: str, reason: str) -> None:
    _para(
        doc,
        f"[NOT_GENERATED_BY_RUN: {section_id} - {reason}]",
        italic=True,
        size=9,
        space_after=4,
    )


def _role_block(
    doc: Document,
    narr_snap: dict[str, Any],
    bullet_snap: dict[str, Any],
    header_key: str,
    *,
    base_header: dict[str, Any] | None = None,
    missing_label: str = "role",
) -> None:
    hdr = base_header or narr_snap.get(header_key) or bullet_snap.get(header_key) or {}
    if not hdr:
        _gap(doc, missing_label, "missing_generated_role_section")
        return
    header_lines = exp_header(hdr)
    if header_lines:
        _para(doc, header_lines[0], bold=True, size=11, space_after=0)
        if len(header_lines) > 1:
            _para(doc, header_lines[1], italic=True, size=9, space_after=2)
    narrative = str(narr_snap.get("narrative_sentence") or "").strip()
    if narrative:
        _para(doc, narrative, italic=True, size=10)
    bullet_lines = bullets_from_list(bullet_snap.get("bullets") or [])
    for line in bullet_lines:
        _para(doc, line, size=10)
    if not narrative and not bullet_lines:
        _gap(doc, missing_label, "generated narrative and bullets were not produced in this run")


def export(final_resume_path: Path, out_path: Path | None = None) -> Path:
    blob = json.loads(final_resume_path.read_text(encoding="utf-8"))
    identity = blob.get("candidate_identity") or {}
    sections = [s for s in (blob.get("sections") or []) if isinstance(s, dict)]
    by_id = {s.get("section_id"): s for s in sections}

    # Refuse to emit a hollow DOCX: the spine schema must carry a candidate name
    # and at least one dict section. A flat ``outputs/`` export (candidate_name
    # top-level, ``sections`` as a dict) lands here with empty ``sections`` and
    # would otherwise render only the "Candidate" + "Professional Experience"
    # placeholders. Fail loudly instead.
    if not str(identity.get("candidate_name") or "").strip() or not sections:
        raise ValueError(
            f"final_resume.json at {final_resume_path} is not spine-shaped "
            "(missing candidate_identity.candidate_name or a list of sections) — "
            "refusing to emit a hollow DOCX. Point at "
            "final_resume_assembly/final_resume.json."
        )

    def snap(sid: str) -> dict[str, Any]:
        sec = by_id.get(sid) or {}
        s = sec.get("l2_output_snapshot")
        return s if isinstance(s, dict) else {}

    doc = Document()
    for section in doc.sections:
        section.top_margin = section.bottom_margin = Pt(36)
        section.left_margin = section.right_margin = Pt(48)

    name = str(identity.get("candidate_name") or "Candidate").strip()
    _para(doc, name, bold=True, size=18, center=True, space_after=0)
    cl = contact_line(identity.get("header_contact") or {})
    if cl:
        _para(doc, cl, size=9, center=True, space_after=4)

    _heading(doc, "Headline")
    headline = str(snap("headline").get("headline_line") or "").strip()
    if headline:
        _para(doc, headline, bold=True, size=11, center=True, space_after=6)
    else:
        _gap(doc, "headline", "missing_or_empty_headline")

    _heading(doc, "Executive Summary")
    exec_text = str(snap("executive_summary").get("resume_display_text") or "").strip()
    if exec_text:
        _para(doc, exec_text, size=10, space_after=4)
    else:
        _gap(doc, "executive_summary", "missing_or_empty_executive_summary")

    comp_cats = snap("competencies").get("competencies") or []
    comp_display = str(snap("competencies").get("resume_display_text") or "").strip()
    _heading(doc, ENGINEERING_PLATFORM_COMPETENCIES_HEADING)
    if comp_cats:
        for cat in comp_cats:
            if not isinstance(cat, dict):
                continue
            label = str(cat.get("category_label") or "Capabilities").strip()
            terms = []
            for t in cat.get("terms") or []:
                txt = str(t.get("text") if isinstance(t, dict) else t).strip()
                if txt:
                    terms.append(txt)
            if terms:
                para = doc.add_paragraph()
                run = para.add_run(f"{label}: ")
                run.bold = True
                run.font.size = Pt(10)
                run2 = para.add_run(", ".join(terms))
                run2.font.size = Pt(10)
                para.paragraph_format.space_after = Pt(2)
    elif comp_display:
        _para(doc, comp_display, size=10)
    else:
        _gap(doc, "competencies", "missing_competencies_or_empty_graph_skills_signal")

    base_headers = base_role_headers_from_final_resume(blob)

    _heading(doc, "Professional Experience")
    _role_block(
        doc,
        snap("unify_narrative"),
        snap("unify_bullets"),
        "unify_header",
        base_header=base_headers.get("unify"),
        missing_label="unify",
    )
    _role_block(
        doc,
        snap("ibm_narrative"),
        snap("ibm_bullets"),
        "ibm_header",
        base_header=base_headers.get("ibm"),
        missing_label="ibm",
    )
    _role_block(
        doc,
        snap("insurtech_narrative"),
        snap("insurtech_bullets"),
        "insurtech_header",
        base_header=base_headers.get("insurtech"),
        missing_label="insurtech",
    )
    _role_block(
        doc,
        snap("ey_narrative"),
        snap("ey_bullets"),
        "ey_header",
        base_header=base_headers.get("ey"),
        missing_label="ey",
    )

    early = (by_id.get("early_career") or {}).get("copied_text_exact")
    if early:
        obj = json.loads(early)
        hdr = base_headers.get("early_career") or {
            k: obj.get(k) for k in ("employer", "title", "location", "start_date", "end_date", "is_current")
        }
        hdr_lines = exp_header(hdr)
        if hdr_lines:
            _para(doc, hdr_lines[0], bold=True, size=11, space_after=0)
            if len(hdr_lines) > 1:
                _para(doc, hdr_lines[1], italic=True, size=9, space_after=2)
        role_narr = str(obj.get("role_narrative") or "").strip()
        if role_narr:
            _para(doc, role_narr, italic=True, size=10)
        for line in bullets_from_list(obj.get("bullets") or []):
            _para(doc, line, size=10)
    else:
        _gap(doc, "early_career", "missing_locked_copy_section")

    edu_raw = (by_id.get("education") or {}).get("copied_text_exact")
    _heading(doc, "Education")
    if edu_raw:
        for entry in json.loads(edu_raw):
            if isinstance(entry, dict):
                line = format_edu(entry)
                if line:
                    _para(doc, line, size=10)
    else:
        _gap(doc, "education", "missing_locked_education")

    cert_raw = (by_id.get("certifications") or {}).get("copied_text_exact")
    _heading(doc, CERTIFICATIONS_AND_CREDENTIALS_HEADING)
    if cert_raw:
        for entry in json.loads(cert_raw):
            if isinstance(entry, dict):
                line = format_cert(entry)
                if line:
                    _para(doc, line, size=10)
    else:
        _gap(doc, "certifications", "missing_locked_certifications")

    if out_path is None:
        stamp = datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S")
        out_path = final_resume_path.parent / f"AmitAyer_AIG_VP_AgenticAI_{stamp}.docx"
    doc.save(str(out_path))
    return out_path


def main(argv: list[str]) -> int:
    if not argv:
        print("usage: export_final_resume_docx.py <run_dir_or_final_resume.json> [out.docx]")
        return 2
    final_path = _resolve_final_resume(argv[0])
    out = Path(argv[1]) if len(argv) > 1 else None
    written = export(final_path, out)
    print(f"DOCX_WRITTEN={written}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main(sys.argv[1:]))

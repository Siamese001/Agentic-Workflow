"""Export a generated_resume_*.json + canonical master_resume into a polished .docx.

Reads:
  - apps_rg/scripts/generated_resume_<ts>.json  (ranked content from the run)
  - apps_shared/data/master_resume.json         (canonical master record for contact, education, certs)
    Legacy `apps_rg/scripts/your_resume_updated.json` was retired 2026-04-30.

Writes:
  - apps_rg/scripts/AmitAyer_SVPAISolutions_<ts>.docx

Usage:
    python apps_rg/scripts/export_to_docx.py [generated_resume_filename]
"""
from __future__ import annotations

import json
import sys
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

SCRIPT_DIR = Path(__file__).parent


def latest_generated() -> Path:
    files = sorted(SCRIPT_DIR.glob("generated_resume_*.json"))
    if not files:
        raise FileNotFoundError("No generated_resume_*.json found in apps_rg/scripts/")
    return files[-1]


def _add_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    # Visual separator as a second paragraph of underscores — no XML injection
    sep = doc.add_paragraph()
    sep_run = sep.add_run("_" * 90)
    sep_run.font.size = Pt(6)
    sep_run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    sep.paragraph_format.space_before = Pt(0)
    sep.paragraph_format.space_after = Pt(2)


def _add_para(doc: Document, text: str, bold: bool = False, size: int = 10, italic: bool = False, align: int | None = None) -> None:
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(2)


def _add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(2)


def build_doc(generated_path: Path, master_path: Path, out_path: Path) -> None:
    generated = json.loads(generated_path.read_text(encoding="utf-8"))
    master = json.loads(master_path.read_text(encoding="utf-8"))

    doc = Document()
    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    contact = master.get("contact_info", {})
    name = contact.get("name", "")
    headline = master.get("headline", "")

    # Header — name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(name)
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    p.paragraph_format.space_after = Pt(0)

    # Headline
    if headline:
        _add_para(doc, headline, italic=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)

    # Contact line
    contact_parts = [
        contact.get("phone", ""),
        contact.get("email", ""),
        contact.get("linkedin", "").replace("https://www.", "").replace("https://", ""),
        contact.get("github", "").replace("https://www.", "").replace("https://", ""),
    ]
    contact_line = " | ".join(c for c in contact_parts if c)
    _add_para(doc, contact_line, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)

    # Summary
    summary = master.get("summary", "")
    if summary:
        _add_heading(doc, "EXECUTIVE SUMMARY")
        _add_para(doc, summary, size=10)

    # Experience — use ranked bullets from generated_resume; fall back to master
    _add_heading(doc, "PROFESSIONAL EXPERIENCE")
    gen_exp = generated.get("experience", [])
    master_exp = {e["company"]: e for e in master.get("experience", []) if isinstance(e, dict) and "company" in e}

    for role in gen_exp:
        company = role.get("company", "")
        title = role.get("title", "")
        m = master_exp.get(company, {})
        location = m.get("location", "")
        duration = m.get("duration", "")
        description = m.get("description", "")

        # Company | Location
        p = doc.add_paragraph()
        run = p.add_run(f"{company}")
        run.bold = True
        run.font.size = Pt(11)
        if location:
            run2 = p.add_run(f"  |  {location}")
            run2.font.size = Pt(10)
            run2.italic = True
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(0)

        # Title | Duration
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.bold = True
        run.italic = True
        run.font.size = Pt(10)
        if duration:
            run2 = p.add_run(f"  |  {duration}")
            run2.font.size = Pt(10)
        p.paragraph_format.space_after = Pt(2)

        # Description (if present in master)
        if description:
            _add_para(doc, description, size=10)

        # Bullets — use ranked order from generated
        for bullet in role.get("bullets", []):
            text = bullet.get("bullet_text") if isinstance(bullet, dict) else str(bullet)
            if text:
                _add_bullet(doc, text)

    # Education
    education = master.get("education", [])
    if education:
        _add_heading(doc, "EDUCATION")
        for edu in education:
            degree = edu.get("degree", "")
            school = edu.get("school", "")
            year = edu.get("year")
            notes = edu.get("notes", "")
            parts = [degree]
            if school:
                parts.append(school)
            if year:
                parts.append(str(year))
            line = ", ".join(parts)
            if notes:
                line += f"  ({notes})"
            _add_para(doc, line, size=10)

    # Certifications
    certs = master.get("certifications_and_credentials", [])
    if certs:
        _add_heading(doc, "CERTIFICATIONS & CREDENTIALS")
        for c in certs:
            name_c = c.get("name", "")
            year_c = c.get("year", "")
            line = f"{name_c}" + (f", {year_c}" if year_c else "")
            _add_para(doc, line, size=10)

    # Competencies (from master) — render the categorized dict
    comps = master.get("competencies", {})
    if comps:
        _add_heading(doc, "ENGINEERING & PLATFORM COMPETENCIES")
        for category, items in comps.items():
            p = doc.add_paragraph()
            run = p.add_run(f"{category}: ")
            run.bold = True
            run.font.size = Pt(10)
            if isinstance(items, list):
                run2 = p.add_run("; ".join(items))
            else:
                run2 = p.add_run(str(items))
            run2.font.size = Pt(10)
            p.paragraph_format.space_after = Pt(2)

    doc.save(out_path)
    print(f"OK Wrote: {out_path}")


def main(argv: list[str]) -> int:
    if len(argv) > 1:
        gen = SCRIPT_DIR / argv[1] if not Path(argv[1]).is_absolute() else Path(argv[1])
    else:
        gen = latest_generated()
    master = Path(__file__).resolve().parents[2] / "apps_shared" / "data" / "master_resume.json"
    if not master.exists():
        raise FileNotFoundError(
            f"Canonical master_resume.json not found at {master}. "
            "Legacy your_resume_updated.json fallback was removed on 2026-04-30."
        )
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    out = SCRIPT_DIR / f"AmitAyer_SVPAISolutions_{ts}.docx"
    build_doc(gen, master, out)
    return 0


if __name__ == "__main__":
    sys.exit(main(sys.argv))
